from __future__ import annotations

"""
abbreviation_list_inserter.py

Точечная доработка для задачи 2.

Что исправлено:
1. В режиме append_existing_list существующий перечень ищется только
   как НЕПРЕРЫВНЫЙ блок сразу после заголовка раздела.
2. Сбор старых записей прекращается на первом обычном тексте/таблице,
   не похожем на перечень сокращений.
3. Добавлена строгая проверка:
   "похоже ли левая часть на сокращение".
4. Поэтому строки вида:
   - "Адрес Сервера администрирования"
   - "Площадка 1"
   - "Анализ журналов"
   больше не попадают в перечень сокращений.
5. Сохраняются режимы:
   - separate_file
   - insert_end
   - insert_before_marker
   - append_existing_list
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
import regex
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table


@dataclass
class AbbreviationEntry:
    abbreviation: str
    long_form: str


class AbbreviationListInserter:
    def __init__(self) -> None:
        self.section_titles = [
            "Перечень обозначений и сокращений",
            "Обозначения и сокращения",
        ]
        self.inline_entry_pattern = regex.compile(
            r"^\s*(?P<abbr>[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9\- ]{0,40})\s*[–\-—]\s*(?P<long>.+?)\s*\.?\s*$"
        )
        self.header_abbr_keywords = {"сокращение", "обозначение"}
        self.header_long_keywords = {"полная форма", "расшифровка"}

    # -----------------------------------------------------------------
    # Подготовка данных
    # -----------------------------------------------------------------

    def _safe_text(self, value) -> str:
        if value is None:
            return ""
        try:
            if pd.isna(value):
                return ""
        except Exception:
            pass
        return " ".join(str(value).split()).strip()

    def _normalize_entries(self, entries: List[AbbreviationEntry]) -> List[AbbreviationEntry]:
        dedup: Dict[Tuple[str, str], AbbreviationEntry] = {}

        for item in entries:
            abbreviation = self._safe_text(item.abbreviation)
            long_form = self._safe_text(item.long_form)

            if not abbreviation or not long_form:
                continue

            key = (abbreviation.upper(), long_form.lower())
            dedup[key] = AbbreviationEntry(
                abbreviation=abbreviation,
                long_form=long_form
            )

        normalized = list(dedup.values())
        normalized.sort(key=lambda x: (x.abbreviation.upper(), x.long_form.lower()))
        return normalized

    def load_entries_from_dataframe(self, df: pd.DataFrame) -> List[AbbreviationEntry]:
        if df.empty:
            return []

        entries: List[AbbreviationEntry] = []

        if {"abbreviation", "long_form"}.issubset(df.columns):
            for _, row in df.iterrows():
                abbreviation = self._safe_text(row.get("abbreviation", ""))
                long_form = self._safe_text(row.get("long_form", ""))
                entries.append(AbbreviationEntry(abbreviation=abbreviation, long_form=long_form))

        elif {"found_abbreviation", "term"}.issubset(df.columns):
            working_df = df.copy()
            if "abbreviation_found_in_text" in working_df.columns:
                working_df = working_df[working_df["abbreviation_found_in_text"] == True]

            for _, row in working_df.iterrows():
                abbreviation = self._safe_text(row.get("found_abbreviation", ""))
                long_form = self._safe_text(row.get("term", ""))
                entries.append(AbbreviationEntry(abbreviation=abbreviation, long_form=long_form))
        else:
            raise ValueError(
                "Не удалось распознать структуру DataFrame. "
                "Ожидаются колонки abbreviation+long_form или found_abbreviation+term."
            )

        return self._normalize_entries(entries)

    def load_entries_from_file(self, path: str | Path) -> List[AbbreviationEntry]:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path}")

        suffix = path.suffix.lower()
        if suffix == ".csv":
            df = pd.read_csv(path, encoding="utf-8-sig")
        elif suffix in {".xlsx", ".xls"}:
            df = pd.read_excel(path)
        else:
            raise ValueError("Поддерживаются только CSV/XLSX файлы.")

        return self.load_entries_from_dataframe(df)

    # -----------------------------------------------------------------
    # Оформление
    # -----------------------------------------------------------------

    def _apply_basic_style(self, document: Document) -> None:
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

    def _add_title(self, document: Document, title: str):
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["Normal"]
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        paragraph.alignment = 1
        return paragraph

    def _add_entries_table(self, document: Document, entries: List[AbbreviationEntry]):
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Сокращение"
        hdr[1].text = "Полная форма"

        for item in entries:
            row = table.add_row().cells
            row[0].text = item.abbreviation
            row[1].text = item.long_form

        return table

    # -----------------------------------------------------------------
    # XML-хелперы
    # -----------------------------------------------------------------

    def _insert_paragraph_after(self, paragraph, text: str = ""):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = paragraph._parent.add_paragraph()
        new_paragraph._p.getparent().remove(new_paragraph._p)
        new_paragraph._p = new_p
        if text:
            new_paragraph.add_run(text)
        return new_paragraph

    def _move_table_after_paragraph(self, paragraph, table):
        paragraph._p.addnext(table._tbl)

    def _remove_element(self, element) -> None:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    # -----------------------------------------------------------------
    # Создание отдельного документа
    # -----------------------------------------------------------------

    def create_separate_document(
        self,
        entries: List[AbbreviationEntry],
        output_path: str | Path,
        title: str = "Перечень обозначений и сокращений"
    ) -> Path:
        entries = self._normalize_entries(entries)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._apply_basic_style(doc)
        self._add_title(doc, title)
        doc.add_paragraph("")
        self._add_entries_table(doc, entries)
        doc.save(output_path)

        return output_path

    def create_separate_document_next_to_source(
        self,
        source_docx_path: str | Path,
        entries: List[AbbreviationEntry],
        suffix_name: str = "_abbreviation_list"
    ) -> Path:
        source_docx_path = Path(source_docx_path)
        output_path = source_docx_path.with_name(
            f"{source_docx_path.stem}{suffix_name}.docx"
        )
        return self.create_separate_document(entries, output_path)

    # -----------------------------------------------------------------
    # Поиск мест вставки
    # -----------------------------------------------------------------

    def _paragraph_style_name(self, paragraph) -> str:
        try:
            if paragraph.style is not None and paragraph.style.name:
                return str(paragraph.style.name).strip()
        except Exception:
            pass
        return ""

    def _is_toc_style(self, paragraph) -> bool:
        style_name = self._paragraph_style_name(paragraph).lower()
        return "toc" in style_name or "оглав" in style_name

    def _looks_like_toc_line(self, text: str) -> bool:
        text = self._safe_text(text)
        if not text:
            return False
        parts = text.split()
        if len(parts) < 2:
            return False
        return parts[-1].isdigit()

    def _is_heading_like(self, paragraph) -> bool:
        style_name = self._paragraph_style_name(paragraph).lower()
        return "heading" in style_name or "заголов" in style_name

    def _find_heading_paragraph(self, document: Document, title_candidates: List[str]):
        lowered_candidates = [item.lower() for item in title_candidates]
        for paragraph in document.paragraphs:
            text = self._safe_text(paragraph.text).lower()
            if text in lowered_candidates and not self._is_toc_style(paragraph):
                return paragraph
        return None

    def _find_paragraph_containing_text(self, document: Document, marker_text: str):
        marker_text = self._safe_text(marker_text).lower()
        if not marker_text:
            return None

        for paragraph in document.paragraphs:
            text = self._safe_text(paragraph.text).lower()
            if text == marker_text and self._is_heading_like(paragraph) and not self._is_toc_style(paragraph):
                return paragraph

        for paragraph in document.paragraphs:
            text = self._safe_text(paragraph.text).lower()
            if not text:
                continue
            if self._is_toc_style(paragraph) or self._looks_like_toc_line(text):
                continue
            if text == marker_text:
                return paragraph

        for paragraph in document.paragraphs:
            text = self._safe_text(paragraph.text).lower()
            if not text:
                continue
            if self._is_toc_style(paragraph) or self._looks_like_toc_line(text):
                continue
            if marker_text in text:
                return paragraph

        return None

    def _looks_like_abbreviation(self, text: str) -> bool:
        """
        Проверяет, похожа ли строка слева на сокращение.
        Строгий фильтр, чтобы не пропускать:
        - Площадка 1
        - Анализ журналов
        - Адрес Сервера администрирования
        """
        text = self._safe_text(text)
        if not text:
            return False
        if len(text) > 40:
            return False

        # Однотокенные варианты: АСО, ПЛК, KSC, СрЗИ, NAP100
        if regex.fullmatch(r"[A-ZА-ЯЁ][A-ZА-ЯЁ0-9\-]{1,20}", text):
            return True

        if regex.fullmatch(r"[А-ЯЁA-Z][A-Za-zА-Яа-яЁё]{1,15}", text):
            uppercase_count = sum(1 for ch in text if ch.isupper())
            return uppercase_count >= 2

        # Многотокенные варианты: KICS for Nodes
        parts = text.split()
        if 1 < len(parts) <= 4:
            has_abbr_like_token = False
            for part in parts:
                if regex.fullmatch(r"[A-ZА-ЯЁ][A-ZА-ЯЁ0-9\-]{1,20}", part):
                    has_abbr_like_token = True
                    break
                if regex.fullmatch(r"[А-ЯЁA-Z][A-Za-zА-Яа-яЁё]{1,15}", part):
                    uppercase_count = sum(1 for ch in part if ch.isupper())
                    if uppercase_count >= 2:
                        has_abbr_like_token = True
                        break
            return has_abbr_like_token

        return False

    def _is_tbl_element(self, element) -> bool:
        return element is not None and element.tag.endswith("}tbl")

    def _is_paragraph_element(self, element) -> bool:
        return element is not None and element.tag.endswith("}p")

    def _paragraph_text_from_element(self, element) -> str:
        if element is None:
            return ""
        texts = element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        return self._safe_text(" ".join(t.text for t in texts if t.text))

    def _table_header_looks_like_abbreviation_table(self, table: Table) -> bool:
        if len(table.rows) < 1 or len(table.rows[0].cells) < 2:
            return False

        left = self._safe_text(table.rows[0].cells[0].text).lower()
        right = self._safe_text(table.rows[0].cells[1].text).lower()

        left_ok = any(key in left for key in self.header_abbr_keywords)
        right_ok = any(key in right for key in self.header_long_keywords)
        return left_ok and right_ok

    def _table_rows_look_like_abbreviation_entries(self, table: Table) -> bool:
        if len(table.rows) < 2:
            return False

        sample_rows = table.rows[: min(len(table.rows), 6)]
        valid = 0
        total = 0

        for idx, row in enumerate(sample_rows):
            if len(row.cells) < 2:
                continue
            left = self._safe_text(row.cells[0].text)
            right = self._safe_text(row.cells[1].text)

            # пропускаем заголовок
            if idx == 0 and self._table_header_looks_like_abbreviation_table(table):
                continue

            if not left and not right:
                continue

            total += 1
            if self._looks_like_abbreviation(left) and right:
                valid += 1

        return total > 0 and valid >= max(1, total // 2 + total % 2)

    def _table_looks_like_abbreviation_table(self, table: Table) -> bool:
        return self._table_header_looks_like_abbreviation_table(table) or self._table_rows_look_like_abbreviation_entries(table)

    def _collect_entries_from_table(self, table: Table) -> List[AbbreviationEntry]:
        entries: List[AbbreviationEntry] = []
        rows = table.rows
        if not rows:
            return entries

        start_index = 1 if self._table_header_looks_like_abbreviation_table(table) else 0

        for row in rows[start_index:]:
            if len(row.cells) < 2:
                continue
            abbreviation = self._safe_text(row.cells[0].text)
            long_form = self._safe_text(row.cells[1].text)
            if not self._looks_like_abbreviation(abbreviation):
                continue
            if abbreviation and long_form:
                entries.append(AbbreviationEntry(abbreviation=abbreviation, long_form=long_form))

        return self._normalize_entries(entries)

    def _parse_inline_entry(self, text: str) -> Optional[AbbreviationEntry]:
        text = self._safe_text(text)
        if not text:
            return None

        match = self.inline_entry_pattern.match(text)
        if not match:
            return None

        abbreviation = self._safe_text(match.group("abbr"))
        long_form = self._safe_text(match.group("long"))

        if not self._looks_like_abbreviation(abbreviation):
            return None
        if not long_form:
            return None

        return AbbreviationEntry(abbreviation=abbreviation, long_form=long_form)

    def _collect_existing_block_after_heading(self, document: Document, heading_paragraph):
        """
        Ищет существующий перечень как непрерывный блок сразу после заголовка.
        Возвращает:
        - entries
        - nodes_to_remove

        Блок может быть:
        - таблицей сокращений;
        - серией строк "АББР – полная форма";
        - смешанным случаем (таблица + строки),
          но только пока это выглядит как перечень сокращений.
        """
        entries: List[AbbreviationEntry] = []
        nodes_to_remove = []

        node = heading_paragraph._p.getnext()

        # пропускаем пустые абзацы сразу после заголовка
        while node is not None and self._is_paragraph_element(node):
            text = self._paragraph_text_from_element(node)
            if text:
                break
            nodes_to_remove.append(node)
            node = node.getnext()

        started = False

        while node is not None:
            if self._is_tbl_element(node):
                table = Table(node, document._body)
                if not self._table_looks_like_abbreviation_table(table):
                    break

                started = True
                nodes_to_remove.append(node)
                entries.extend(self._collect_entries_from_table(table))
                node = node.getnext()
                continue

            if self._is_paragraph_element(node):
                text = self._paragraph_text_from_element(node)

                if not text:
                    if started:
                        nodes_to_remove.append(node)
                    node = node.getnext()
                    continue

                entry = self._parse_inline_entry(text)
                if entry is None:
                    break

                started = True
                nodes_to_remove.append(node)
                entries.append(entry)
                node = node.getnext()
                continue

            break

        return self._normalize_entries(entries), nodes_to_remove

    # -----------------------------------------------------------------
    # Вставка в существующий документ
    # -----------------------------------------------------------------

    def insert_into_existing_document(
        self,
        source_docx_path: str | Path,
        entries: List[AbbreviationEntry],
        output_path: str | Path,
        mode: str = "end",
        marker_text: Optional[str] = None,
        section_title: str = "Перечень обозначений и сокращений"
    ) -> Path:
        entries = self._normalize_entries(entries)
        source_docx_path = Path(source_docx_path)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if not source_docx_path.exists():
            raise FileNotFoundError(f"Файл не найден: {source_docx_path}")

        doc = Document(source_docx_path)
        self._apply_basic_style(doc)

        if mode == "end":
            self._insert_at_end(doc, entries, section_title)

        elif mode == "before_marker":
            if not marker_text:
                raise ValueError("Для режима before_marker нужно указать marker_text.")
            self._insert_before_marker(doc, entries, marker_text, section_title)

        elif mode == "append_existing_list":
            self._append_to_existing_or_create(doc, entries, section_title)

        else:
            raise ValueError(
                "Неизвестный режим вставки. Поддерживаются: "
                "'end', 'before_marker', 'append_existing_list'."
            )

        doc.save(output_path)
        return output_path

    def _insert_at_end(
        self,
        document: Document,
        entries: List[AbbreviationEntry],
        section_title: str
    ) -> None:
        document.add_page_break()
        self._add_title(document, section_title)
        document.add_paragraph("")
        self._add_entries_table(document, entries)

    def _insert_before_marker(
        self,
        document: Document,
        entries: List[AbbreviationEntry],
        marker_text: str,
        section_title: str
    ) -> None:
        target = self._find_paragraph_containing_text(document, marker_text)
        if target is None:
            raise ValueError(f'Не найден маркер для вставки: "{marker_text}"')

        title_par = target.insert_paragraph_before("")
        title_run = title_par.add_run(section_title)
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(14)
        title_par.alignment = 1

        spacer = target.insert_paragraph_before("")
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Сокращение"
        hdr[1].text = "Полная форма"

        for item in entries:
            row = table.add_row().cells
            row[0].text = item.abbreviation
            row[1].text = item.long_form

        self._move_table_after_paragraph(spacer, table)

    def _append_to_existing_or_create(
        self,
        document: Document,
        entries: List[AbbreviationEntry],
        section_title: str
    ) -> None:
        existing_heading = self._find_heading_paragraph(document, self.section_titles)

        if existing_heading is None:
            self._insert_at_end(document, entries, section_title)
            return

        existing_entries, nodes_to_remove = self._collect_existing_block_after_heading(document, existing_heading)
        merged_entries = self._normalize_entries(existing_entries + entries) if existing_entries else self._normalize_entries(entries)

        # удаляем только непрерывный старый блок перечня сразу после заголовка
        for node in reversed(nodes_to_remove):
            self._remove_element(node)

        spacer = self._insert_paragraph_after(existing_heading, "")
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Сокращение"
        hdr[1].text = "Полная форма"

        for item in merged_entries:
            row = table.add_row().cells
            row[0].text = item.abbreviation
            row[1].text = item.long_form

        self._move_table_after_paragraph(spacer, table)

    # -----------------------------------------------------------------
    # Высокоуровневый запуск
    # -----------------------------------------------------------------

    def run(
        self,
        input_data_path: str | Path,
        source_docx_path: str | Path,
        mode: str = "separate_file",
        output_path: Optional[str | Path] = None,
        marker_text: Optional[str] = None
    ) -> Dict[str, Path]:
        entries = self.load_entries_from_file(input_data_path)
        source_docx_path = Path(source_docx_path)

        if mode == "separate_file":
            if output_path is None:
                created = self.create_separate_document_next_to_source(source_docx_path, entries)
            else:
                created = self.create_separate_document(entries, output_path)
            return {"output_docx": created}

        if output_path is None:
            output_path = source_docx_path.with_name(f"{source_docx_path.stem}_with_abbreviation_list.docx")

        if mode == "insert_end":
            created = self.insert_into_existing_document(
                source_docx_path=source_docx_path,
                entries=entries,
                output_path=output_path,
                mode="end"
            )
            return {"output_docx": created}

        if mode == "insert_before_marker":
            created = self.insert_into_existing_document(
                source_docx_path=source_docx_path,
                entries=entries,
                output_path=output_path,
                mode="before_marker",
                marker_text=marker_text
            )
            return {"output_docx": created}

        if mode == "append_existing_list":
            created = self.insert_into_existing_document(
                source_docx_path=source_docx_path,
                entries=entries,
                output_path=output_path,
                mode="append_existing_list"
            )
            return {"output_docx": created}

        raise ValueError(
            "Неизвестный режим run(). Поддерживаются: "
            "'separate_file', 'insert_end', 'insert_before_marker', 'append_existing_list'."
        )


if __name__ == "__main__":
    inserter = AbbreviationListInserter()

    input_data_path = "result_all/stage2/existing_abbreviations.csv"
    source_docx_path = "test_reduction_input.docx"

    result = inserter.run(
        input_data_path=input_data_path,
        source_docx_path=source_docx_path,
        mode="append_existing_list"
    )

    print("=" * 72)
    print("ПЕРЕЧЕНЬ СОКРАЩЕНИЙ СФОРМИРОВАН")
    print("=" * 72)
    for name, path in result.items():
        print(f"{name}: {path}")
