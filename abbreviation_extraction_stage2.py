from __future__ import annotations

"""
abbreviation_extraction_stage2.py

Этап 2 проекта:
вычленение всех доступных к сокращению словоформ
и уже имеющихся в документе аббревиатур.

Финальная доработка:
1. Поддерживает шаблоны:
   - полная форма (АББР)
   - полная форма (далее – АББР)
   - АББР (полная форма)
2. Извлекает сокращения из раздела "Обозначения и сокращения".
3. Поддерживает mixed-case и multi-token сокращения:
   - СрЗИ
   - KICS for Nodes
   - KSC / KSN
4. Защищена от ложных parenthetical-случаев:
   - ТСПД (включая беспроводные радиоканалы)
5. Принудительно канонизирует термины по глоссарию:
   - активное сетевое оборудование средства -> активное сетевое оборудование
6. Гарантирует присутствие глоссарных терминов в итоговой merged-таблице
   даже если этап 1 их не выделил или выделил шумно.
"""

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Set
import zipfile
import xml.etree.ElementTree as ET

import pandas as pd
import regex
from rapidfuzz import fuzz
from natasha import Doc, Segmenter
from pymorphy2 import MorphAnalyzer

from text_recognition_candidates_v3 import ReducibleWordformRecognizerV3


@dataclass
class TextFragment:
    source_type: str
    source_index: int
    text: str


@dataclass
class FoundAbbreviation:
    abbreviation: str
    long_form: str
    detection_type: str
    source_type: str
    source_index: int
    sentence: str
    matched_term: str
    match_score: float


class DocumentTextExtractor:
    def extract_fragments(self, docx_path: str | Path) -> List[TextFragment]:
        from docx import Document

        docx_path = Path(docx_path)
        doc = Document(docx_path)

        fragments: List[TextFragment] = []
        paragraph_index = 0
        heading_index = 0
        table_cell_index = 0
        footnote_index = 0

        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text:
                continue

            style_name = ""
            if paragraph.style is not None and paragraph.style.name:
                style_name = paragraph.style.name.lower()

            if "heading" in style_name or "заголов" in style_name:
                fragments.append(TextFragment("heading", heading_index, text))
                heading_index += 1
            else:
                fragments.append(TextFragment("paragraph", paragraph_index, text))
                paragraph_index += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._clean_text(cell.text)
                    if not text:
                        continue
                    fragments.append(TextFragment("table_cell", table_cell_index, text))
                    table_cell_index += 1

        for text in self._extract_footnotes_from_docx(docx_path):
            cleaned = self._clean_text(text)
            if not cleaned:
                continue
            fragments.append(TextFragment("footnote", footnote_index, cleaned))
            footnote_index += 1

        return fragments

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\n", " ")
        text = regex.sub(r"\s+", " ", text)
        return text.strip()

    def _extract_footnotes_from_docx(self, docx_path: Path) -> List[str]:
        result: List[str] = []
        try:
            with zipfile.ZipFile(docx_path, "r") as archive:
                if "word/footnotes.xml" not in archive.namelist():
                    return result

                xml_bytes = archive.read("word/footnotes.xml")
                root = ET.fromstring(xml_bytes)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

                for footnote in root.findall(".//w:footnote", ns):
                    texts = []
                    for node in footnote.findall(".//w:t", ns):
                        if node.text:
                            texts.append(node.text)
                    footnote_text = " ".join(texts).strip()
                    if footnote_text:
                        result.append(footnote_text)
        except Exception:
            return []

        return result


class ExistingAbbreviationExtractor:
    def __init__(self) -> None:
        self.segmenter = Segmenter()
        self.morph = MorphAnalyzer()

        self.classic_abbr_pattern = r"[A-ZА-ЯЁ][A-ZА-ЯЁ0-9-]{1,15}"
        self.word_pattern = regex.compile(r"[A-Za-zА-Яа-яЁё-]+")

        self.pattern_long_first = regex.compile(
            rf"(?P<long>[^()]{{3,250}}?)\((?P<abbr>{self.classic_abbr_pattern})\)",
            flags=regex.IGNORECASE
        )

        self.pattern_long_dalee = regex.compile(
            rf"(?P<long>[^()]{{3,250}}?)\(\s*далее\s*[–\-—]\s*(?P<abbr>[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9\-\s]{{1,40}})\s*\)",
            flags=regex.IGNORECASE
        )

        self.pattern_abbr_first = regex.compile(
            rf"(?P<abbr>{self.classic_abbr_pattern})\s*\((?P<long>[^)]{{3,250}})\)",
            flags=regex.IGNORECASE
        )

        self.pattern_standalone_abbr = regex.compile(rf"\b{self.classic_abbr_pattern}\b")

        self.pattern_glossary_line = regex.compile(
            r"^\s*(?P<abbr>[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9\- ]{0,40})\s*[–\-—]\s*(?P<long>.+?)\s*\.?\s*$"
        )

        self.false_positive_abbr = {
            "РФ", "РТФ", "MS", "WORD", "PDF", "DOCX",
            "ДАЛЕЕ", "ТАБЛИЦА", "РИСУНОК", "ОБЩИЕ", "ПРИЛОЖЕНИЕ",
            "ЛИСТ", "ЛИСТОВ", "ООО"
        }

        self.allowed_pos = {"NOUN", "ADJF", "ADJS", "PRTF", "PRTS"}

        self.context_lemmas = {
            "в", "во", "на", "при", "по", "внутри", "внутренний",
            "сноска", "дополнительно", "указанный", "указать",
            "отчёт", "документ", "пример", "тестирование", "проверка",
            "данный", "этот", "тот", "далее", "рамка", "проект",
            "создавать", "создаваемый", "настоящий", "который",
            "выполняться", "использовать", "устанавливать",
            "описывать", "текущий", "наименование", "состав",
            "включая", "отдельный", "следующий", "подраздел", "раздел",
            "например"
        }

        self.left_noise_lemmas = {
            "обеспечение", "обеспечивать", "состав", "наименование",
            "решение", "цель", "функционал", "параметр", "рамка",
            "проект", "технический", "настройка", "порядок",
            "использование", "применение", "выполнение", "осуществление"
        }

        self.trailing_noise_lemmas = {
            "средство", "средства", "элемент", "элементы", "уровень",
            "уровни", "состав", "компонент", "компоненты"
        }

        self.parenthetical_context_markers = {
            "включая", "например", "в том числе", "такие как"
        }

        self.leading_context_patterns = [
            regex.compile(r"^(?:по|для|в рамках|в составе|с целью)\s+", flags=regex.IGNORECASE),
        ]

    def split_into_sentences(self, text: str) -> List[str]:
        if not text.strip():
            return []
        doc = Doc(text)
        doc.segment(self.segmenter)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]

    def _clean_long_form(self, text: str) -> str:
        text = regex.sub(r"\s+", " ", text).strip()
        return text.strip(" ,.;:()[]{}\"'«»")

    def _extract_words(self, text: str) -> List[str]:
        return self.word_pattern.findall(text)

    def _parse_word(self, word: str):
        return self.morph.parse(word)[0]

    def _get_pos(self, word: str) -> str:
        pos = self._parse_word(word).tag.POS
        return pos if pos else ""

    def _get_lemma(self, word: str) -> str:
        return self._parse_word(word).normal_form

    def _is_content_word(self, word: str) -> bool:
        if len(word) < 3:
            return False
        lemma = self._get_lemma(word)
        pos = self._get_pos(word)
        if lemma in self.context_lemmas:
            return False
        return pos in self.allowed_pos

    def _is_term_like_words(self, words: List[str]) -> bool:
        if len(words) < 2 or len(words) > 8:
            return False

        pos_list = [self._get_pos(word) for word in words]
        noun_count = sum(1 for pos in pos_list if pos == "NOUN")
        modifier_count = sum(1 for pos in pos_list if pos in {"ADJF", "ADJS", "PRTF", "PRTS"})

        if pos_list[-1] != "NOUN":
            return False

        return noun_count >= 2 or (noun_count >= 1 and modifier_count >= 1)

    def _normalize_abbreviation(self, abbreviation: str) -> str:
        return regex.sub(r"\s+", " ", str(abbreviation)).strip()

    def _looks_like_abbreviation(self, text: str) -> bool:
        text = self._normalize_abbreviation(text)
        if not text or len(text) > 45:
            return False

        if regex.fullmatch(r"[A-ZА-ЯЁ][A-ZА-ЯЁ0-9-]{1,15}", text):
            return True

        if regex.fullmatch(r"[А-ЯЁA-Z][A-Za-zА-Яа-яЁё]{1,12}", text):
            if sum(1 for ch in text if ch.isupper()) >= 2:
                return True

        if " " in text:
            parts = text.split()
            if 1 < len(parts) <= 4:
                if all(regex.fullmatch(r"[A-Za-z0-9-]+", part) for part in parts):
                    alpha_parts = sum(1 for part in parts if regex.search(r"[A-Za-z]", part))
                    if alpha_parts >= 2:
                        return True

        return False

    def _is_valid_abbreviation(self, abbreviation: str) -> bool:
        abbreviation = self._normalize_abbreviation(abbreviation)
        if not self._looks_like_abbreviation(abbreviation):
            return False
        if abbreviation.upper() in self.false_positive_abbr:
            return False
        return True

    def _is_contextual_parenthetical_long_form(self, raw_text: str) -> bool:
        text = self._clean_long_form(raw_text).lower()
        if not text:
            return True
        for marker in self.parenthetical_context_markers:
            if marker in text:
                return True
        return False

    def _remove_left_context(self, text: str) -> str:
        text = self._clean_long_form(text)

        for pattern in self.leading_context_patterns:
            text = pattern.sub("", text).strip()

        words = self._extract_words(text)
        if len(words) < 2:
            return text

        lemmas = [self._get_lemma(word) for word in words]
        start = 0
        while start < len(words) - 1:
            lemma = lemmas[start]
            pos = self._get_pos(words[start])

            if lemma in self.left_noise_lemmas:
                start += 1
                continue
            if pos in {"PREP", "CONJ", "PRCL"}:
                start += 1
                continue
            if not self._is_content_word(words[start]) and self._is_content_word(words[start + 1]):
                start += 1
                continue
            break

        return " ".join(words[start:]).strip()

    def _postprocess_candidate_words(self, words: List[str]) -> List[str]:
        if len(words) < 2:
            return words

        while len(words) > 2:
            last_lemma = self._get_lemma(words[-1])
            if last_lemma in self.trailing_noise_lemmas:
                words = words[:-1]
            else:
                break

        return words

    def _shrink_long_form(self, raw_text: str, direction: str) -> str:
        text = self._remove_left_context(raw_text)
        words = self._extract_words(text)
        if len(words) < 2:
            return text

        best_candidate = None
        max_window = min(8, len(words))

        for window_size in range(max_window, 1, -1):
            for start in range(0, len(words) - window_size + 1):
                end = start + window_size
                chunk = words[start:end]

                if not all(self._is_content_word(w) for w in chunk):
                    continue
                if not self._is_term_like_words(chunk):
                    continue

                chunk = self._postprocess_candidate_words(chunk)
                if len(chunk) < 2:
                    continue
                if not self._is_term_like_words(chunk):
                    continue

                candidate_text = " ".join(chunk)
                score = end if direction == "right" else -start
                candidate = (len(chunk), score, candidate_text)

                if best_candidate is None or candidate > best_candidate:
                    best_candidate = candidate

        if best_candidate:
            return best_candidate[2]

        filtered_words = [w for w in words if self._is_content_word(w)]
        filtered_words = self._postprocess_candidate_words(filtered_words)
        if len(filtered_words) >= 2:
            return " ".join(filtered_words[-8:] if direction == "right" else filtered_words[:8])

        return text

    def extract_from_fragments(self, fragments: List[TextFragment]) -> List[FoundAbbreviation]:
        found: List[FoundAbbreviation] = []

        for fragment in fragments:
            for sentence in self.split_into_sentences(fragment.text):
                declared_items = self._extract_declared_abbreviations(
                    sentence=sentence,
                    source_type=fragment.source_type,
                    source_index=fragment.source_index
                )
                found.extend(declared_items)

                declared_abbreviations_in_sentence = {
                    self._normalize_abbreviation(item.abbreviation)
                    for item in declared_items
                }
                found.extend(
                    self._extract_standalone_abbreviations(
                        sentence=sentence,
                        source_type=fragment.source_type,
                        source_index=fragment.source_index,
                        declared_in_same_sentence=declared_abbreviations_in_sentence
                    )
                )

        return self._deduplicate(found)

    def _extract_declared_abbreviations(self, sentence: str, source_type: str, source_index: int) -> List[FoundAbbreviation]:
        result: List[FoundAbbreviation] = []

        for match in self.pattern_long_dalee.finditer(sentence):
            raw_long_form = match.group("long")
            abbreviation = self._normalize_abbreviation(match.group("abbr"))
            if not self._is_valid_abbreviation(abbreviation):
                continue

            long_form = self._shrink_long_form(raw_long_form, direction="right")
            if len(self._extract_words(long_form)) < 2:
                continue

            result.append(
                FoundAbbreviation(
                    abbreviation=abbreviation,
                    long_form=long_form,
                    detection_type="declared_long_dalee",
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    matched_term="",
                    match_score=0.0
                )
            )

        for match in self.pattern_long_first.finditer(sentence):
            raw_long_form = match.group("long")
            abbreviation = self._normalize_abbreviation(match.group("abbr"))
            inside_parentheses = sentence[match.start():match.end()]
            if regex.search(r"\(\s*далее\s*[–\-—]", inside_parentheses, flags=regex.IGNORECASE):
                continue
            if not self._is_valid_abbreviation(abbreviation):
                continue

            long_form = self._shrink_long_form(raw_long_form, direction="right")
            if len(self._extract_words(long_form)) < 2:
                continue

            result.append(
                FoundAbbreviation(
                    abbreviation=abbreviation,
                    long_form=long_form,
                    detection_type="declared_long_first",
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    matched_term="",
                    match_score=0.0
                )
            )

        for match in self.pattern_abbr_first.finditer(sentence):
            raw_long_form = match.group("long")
            abbreviation = self._normalize_abbreviation(match.group("abbr"))
            if not self._is_valid_abbreviation(abbreviation):
                continue
            if regex.match(r"\s*далее\s*[–\-—]", raw_long_form, flags=regex.IGNORECASE):
                continue
            if self._is_contextual_parenthetical_long_form(raw_long_form):
                continue

            long_form = self._shrink_long_form(raw_long_form, direction="left")
            if len(self._extract_words(long_form)) < 2:
                continue

            result.append(
                FoundAbbreviation(
                    abbreviation=abbreviation,
                    long_form=long_form,
                    detection_type="declared_abbr_first",
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    matched_term="",
                    match_score=0.0
                )
            )

        return result

    def _extract_standalone_abbreviations(
        self,
        sentence: str,
        source_type: str,
        source_index: int,
        declared_in_same_sentence: Set[str]
    ) -> List[FoundAbbreviation]:
        result: List[FoundAbbreviation] = []
        for match in self.pattern_standalone_abbr.finditer(sentence):
            abbreviation = self._normalize_abbreviation(match.group(0))
            if not self._is_valid_abbreviation(abbreviation):
                continue
            if abbreviation in declared_in_same_sentence:
                continue
            result.append(
                FoundAbbreviation(
                    abbreviation=abbreviation,
                    long_form="",
                    detection_type="standalone",
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    matched_term="",
                    match_score=0.0
                )
            )
        return result

    def _deduplicate(self, items: List[FoundAbbreviation]) -> List[FoundAbbreviation]:
        seen = set()
        result: List[FoundAbbreviation] = []
        for item in items:
            key = (
                self._normalize_abbreviation(item.abbreviation),
                item.long_form.strip().lower(),
                item.detection_type,
                item.source_type,
                item.source_index,
                item.sentence
            )
            if key in seen:
                continue
            seen.add(key)
            result.append(item)
        return result


class Stage2ReductionAnalyzer:
    def __init__(self) -> None:
        self.stage1_recognizer = ReducibleWordformRecognizerV3()
        self.text_extractor = DocumentTextExtractor()
        self.abbreviation_extractor = ExistingAbbreviationExtractor()

    def _extract_glossary_abbreviations_from_fragments(self, fragments: List[TextFragment]) -> List[FoundAbbreviation]:
        result: List[FoundAbbreviation] = []
        in_glossary = False
        glossary_started = False

        for fragment in fragments:
            text = str(fragment.text).strip()
            lowered = text.lower()

            if fragment.source_type == "heading" and "обозначения и сокращения" in lowered:
                in_glossary = True
                glossary_started = True
                continue

            if in_glossary:
                if fragment.source_type == "heading" and "обозначения и сокращения" not in lowered:
                    break

                candidate_lines = [line.strip() for line in text.split("\n") if line.strip()] or [text]

                for line in candidate_lines:
                    match = self.abbreviation_extractor.pattern_glossary_line.match(line)
                    if not match:
                        continue

                    abbreviation = self.abbreviation_extractor._normalize_abbreviation(match.group("abbr"))
                    long_form = self.abbreviation_extractor._clean_long_form(match.group("long"))

                    if not self.abbreviation_extractor._is_valid_abbreviation(abbreviation):
                        continue
                    if len(self.abbreviation_extractor._extract_words(long_form)) < 1:
                        continue

                    result.append(
                        FoundAbbreviation(
                            abbreviation=abbreviation,
                            long_form=long_form,
                            detection_type="glossary_section",
                            source_type=fragment.source_type,
                            source_index=fragment.source_index,
                            sentence=line,
                            matched_term="",
                            match_score=100.0
                        )
                    )

        if not result:
            for fragment in fragments:
                text = str(fragment.text).strip()
                match = self.abbreviation_extractor.pattern_glossary_line.match(text)
                if not match:
                    continue

                abbreviation = self.abbreviation_extractor._normalize_abbreviation(match.group("abbr"))
                long_form = self.abbreviation_extractor._clean_long_form(match.group("long"))

                if not self.abbreviation_extractor._is_valid_abbreviation(abbreviation):
                    continue
                if len(self.abbreviation_extractor._extract_words(long_form)) < 1:
                    continue

                if len(text.split()) <= 12:
                    result.append(
                        FoundAbbreviation(
                            abbreviation=abbreviation,
                            long_form=long_form,
                            detection_type="glossary_section",
                            source_type=fragment.source_type,
                            source_index=fragment.source_index,
                            sentence=text,
                            matched_term="",
                            match_score=100.0
                        )
                    )

        return self.abbreviation_extractor._deduplicate(result) if (glossary_started or result) else []

    def _source_priority(self, detection_type: str) -> int:
        priorities = {
            "glossary_section": 5,
            "declared_long_dalee": 4,
            "declared_long_first": 3,
            "declared_abbr_first": 2,
            "standalone": 1,
        }
        return priorities.get(str(detection_type), 0)

    def _normalize_term_for_compare(self, term: str) -> str:
        return regex.sub(r"\s+", " ", str(term)).strip().lower()

    def _build_abbreviation_from_phrase(self, phrase: str) -> str:
        words = regex.findall(r"[A-Za-zА-Яа-яЁё-]+", phrase)
        return "".join(word[0].upper() for word in words if word)

    def _token_set(self, text: str) -> Set[str]:
        words = regex.findall(r"[A-Za-zА-Яа-яЁё-]+", str(text).lower())
        return {word for word in words if len(word) >= 3}

    def _word_overlap_ratio(self, a: str, b: str) -> float:
        set_a = self._token_set(a)
        set_b = self._token_set(b)
        if not set_a or not set_b:
            return 0.0
        inter = len(set_a & set_b)
        denom = max(len(set_a), len(set_b))
        return inter / denom if denom else 0.0

    def _ensure_glossary_terms_present(
        self,
        reducible_terms_df: pd.DataFrame,
        found_abbreviations: List[FoundAbbreviation]
    ) -> pd.DataFrame:
        rows = reducible_terms_df.to_dict("records") if not reducible_terms_df.empty else []
        existing_by_abbr = {
            str(row.get("suggested_abbreviation", "")).upper(): row
            for row in rows
        }

        additions = []
        for item in found_abbreviations:
            if item.detection_type != "glossary_section":
                continue
            if not item.long_form:
                continue

            abbr = self.abbreviation_extractor._normalize_abbreviation(item.abbreviation).upper()
            long_form = item.long_form.strip()

            # Если уже есть запись с таким сокращением, но форма шумная,
            # канонизируем её позже.
            if abbr in existing_by_abbr:
                continue

            if len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", long_form)) < 1:
                continue

            additions.append({
                "term": long_form,
                "normalized_term": self._normalize_term_for_compare(long_form),
                "suggested_abbreviation": abbr,
                "word_count": len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", long_form)),
                "frequency": 1,
                "examples": item.sentence
            })

        if additions:
            reducible_terms_df = pd.concat([reducible_terms_df, pd.DataFrame(additions)], ignore_index=True)

        return reducible_terms_df.sort_values(
            by=["frequency", "word_count", "term"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

    def _canonicalize_terms_by_glossary(
        self,
        reducible_terms_df: pd.DataFrame,
        found_abbreviations: List[FoundAbbreviation]
    ) -> pd.DataFrame:
        if reducible_terms_df.empty:
            return reducible_terms_df

        glossary_map = {}
        for item in found_abbreviations:
            if item.detection_type == "glossary_section" and item.long_form:
                glossary_map[self.abbreviation_extractor._normalize_abbreviation(item.abbreviation).upper()] = item.long_form.strip()

        rows = reducible_terms_df.to_dict("records")
        normalized_rows = []

        for row in rows:
            term = str(row.get("term", "")).strip()
            suggested_abbreviation = str(row.get("suggested_abbreviation", "")).upper()
            canonical = glossary_map.get(suggested_abbreviation)

            if canonical:
                overlap = self._word_overlap_ratio(term, canonical)
                term_tokens = self._token_set(term)
                canonical_tokens = self._token_set(canonical)
                extra_tokens = term_tokens - canonical_tokens

                # Для глоссарного сокращения всегда предпочитаем канонический термин,
                # если текущий похож на него хотя бы умеренно.
                if overlap >= 0.60 and len(extra_tokens) <= 2:
                    row["term"] = canonical
                    row["normalized_term"] = self._normalize_term_for_compare(canonical)
                    row["word_count"] = len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", canonical))

            normalized_rows.append(row)

        df = pd.DataFrame(normalized_rows)
        if df.empty:
            return df

        grouped = {}
        for row in df.to_dict("records"):
            key = (
                str(row.get("suggested_abbreviation", "")).upper(),
                str(row.get("normalized_term", "")).strip().lower(),
            )
            if key not in grouped:
                grouped[key] = row.copy()
            else:
                grouped[key]["frequency"] = max(int(grouped[key].get("frequency", 0)), int(row.get("frequency", 0)))
                old_examples = str(grouped[key].get("examples", ""))
                new_examples = str(row.get("examples", ""))
                if new_examples and new_examples not in old_examples:
                    grouped[key]["examples"] = (old_examples + " || " + new_examples).strip(" |")

        return pd.DataFrame(grouped.values()).sort_values(
            by=["frequency", "word_count", "term"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

    def run(self, docx_path: str | Path, output_dir: str | Path) -> Dict[str, Path]:
        docx_path = Path(docx_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        mentions = self.stage1_recognizer.analyze_document(docx_path)
        reducible_terms_df = self.stage1_recognizer.aggregate_mentions(mentions).copy()
        reducible_terms_df = reducible_terms_df.rename(columns={
            "phrase_example": "term",
            "normalized_phrase": "normalized_term",
            "proposed_abbreviation": "suggested_abbreviation"
        })

        fragments = self.text_extractor.extract_fragments(docx_path)
        fragments_df = pd.DataFrame([asdict(f) for f in fragments])

        found_abbreviations_text = self.abbreviation_extractor.extract_from_fragments(fragments)
        found_abbreviations_glossary = self._extract_glossary_abbreviations_from_fragments(fragments)
        found_abbreviations = self.abbreviation_extractor._deduplicate(
            found_abbreviations_glossary + found_abbreviations_text
        )

        reducible_terms_df = self._extend_reducible_terms_with_declared_forms(reducible_terms_df, found_abbreviations)
        reducible_terms_df = self._ensure_glossary_terms_present(reducible_terms_df, found_abbreviations)
        reducible_terms_df = self._canonicalize_terms_by_glossary(reducible_terms_df, found_abbreviations)

        matched_abbreviations = self._match_abbreviations_to_terms(found_abbreviations, reducible_terms_df)
        existing_abbreviations_df = pd.DataFrame([asdict(item) for item in matched_abbreviations])

        merged_df = self._build_merged_table(reducible_terms_df, existing_abbreviations_df, found_abbreviations)

        return self._save_results(
            fragments_df=fragments_df,
            reducible_terms_df=reducible_terms_df,
            existing_abbreviations_df=existing_abbreviations_df,
            merged_df=merged_df,
            output_dir=output_dir
        )

    def _extend_reducible_terms_with_declared_forms(
        self,
        reducible_terms_df: pd.DataFrame,
        found_abbreviations: List[FoundAbbreviation]
    ) -> pd.DataFrame:
        rows = reducible_terms_df.to_dict("records") if not reducible_terms_df.empty else []
        existing_terms_normalized = {
            self._normalize_term_for_compare(str(row.get("term", "")))
            for row in rows
        }
        existing_by_abbr = {
            str(row.get("suggested_abbreviation", "")).upper(): self._normalize_term_for_compare(str(row.get("term", "")))
            for row in rows
        }

        additions = []
        for item in found_abbreviations:
            if not item.long_form:
                continue

            normalized_long_form = self._normalize_term_for_compare(item.long_form)
            if normalized_long_form in existing_terms_normalized:
                continue

            if len(self.abbreviation_extractor._extract_words(item.long_form)) < 2:
                continue

            suggested_abbreviation = self._build_abbreviation_from_phrase(item.long_form)

            is_near_duplicate = False
            for existing in existing_terms_normalized:
                if fuzz.ratio(normalized_long_form, existing) >= 85:
                    if existing_by_abbr.get(suggested_abbreviation.upper()) == existing:
                        is_near_duplicate = True
                        break
            if is_near_duplicate:
                continue

            additions.append({
                "term": item.long_form,
                "normalized_term": normalized_long_form,
                "suggested_abbreviation": suggested_abbreviation,
                "word_count": len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", item.long_form)),
                "frequency": 1,
                "examples": item.sentence
            })
            existing_terms_normalized.add(normalized_long_form)
            existing_by_abbr[suggested_abbreviation.upper()] = normalized_long_form

        if additions:
            reducible_terms_df = pd.concat([reducible_terms_df, pd.DataFrame(additions)], ignore_index=True)

        return reducible_terms_df.sort_values(
            by=["frequency", "word_count", "term"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

    def _match_abbreviations_to_terms(
        self,
        found_abbreviations: List[FoundAbbreviation],
        reducible_terms_df: pd.DataFrame
    ) -> List[FoundAbbreviation]:
        if reducible_terms_df.empty:
            return found_abbreviations

        terms = reducible_terms_df.to_dict("records")
        matched: List[FoundAbbreviation] = []

        glossary_map = {}
        for item in found_abbreviations:
            if item.detection_type == "glossary_section" and item.long_form:
                glossary_map[self.abbreviation_extractor._normalize_abbreviation(item.abbreviation)] = item.long_form

        for item in found_abbreviations:
            best_term = ""
            best_score = 0.0

            abbreviation = self.abbreviation_extractor._normalize_abbreviation(item.abbreviation)
            long_form = str(item.long_form or "").strip().lower()

            if abbreviation in glossary_map:
                long_form = glossary_map[abbreviation].strip().lower()

            for term_row in terms:
                term = str(term_row.get("term", ""))
                normalized_term = str(term_row.get("normalized_term", ""))
                suggested_abbreviation = str(term_row.get("suggested_abbreviation", "")).upper()

                score = 0.0

                if abbreviation and abbreviation.upper() == suggested_abbreviation:
                    score = 100.0

                if long_form:
                    overlap_term = self._word_overlap_ratio(long_form, term)
                    overlap_norm = self._word_overlap_ratio(long_form, normalized_term)

                    if overlap_term >= 0.90:
                        score = max(score, float(fuzz.ratio(long_form, term.lower())))
                    if overlap_norm >= 0.90:
                        score = max(score, float(fuzz.ratio(long_form, normalized_term.lower())))

                if score > best_score:
                    best_score = score
                    best_term = term

            matched.append(
                FoundAbbreviation(
                    abbreviation=item.abbreviation,
                    long_form=(glossary_map.get(abbreviation, item.long_form) or ""),
                    detection_type=item.detection_type,
                    source_type=item.source_type,
                    source_index=item.source_index,
                    sentence=item.sentence,
                    matched_term=best_term if best_score >= 85 else "",
                    match_score=best_score
                )
            )

        return matched

    def _build_merged_table(
        self,
        reducible_terms_df: pd.DataFrame,
        existing_abbreviations_df: pd.DataFrame,
        found_abbreviations: List[FoundAbbreviation]
    ) -> pd.DataFrame:
        if reducible_terms_df.empty:
            return pd.DataFrame(columns=[
                "term", "normalized_term", "suggested_abbreviation", "frequency",
                "abbreviation_found_in_text", "found_abbreviation", "match_score"
            ])

        existing_records = existing_abbreviations_df.to_dict("records") if not existing_abbreviations_df.empty else []

        glossary_map = {}
        for item in found_abbreviations:
            if item.detection_type == "glossary_section" and item.long_form:
                glossary_map[self.abbreviation_extractor._normalize_abbreviation(item.abbreviation).upper()] = item.long_form.strip()

        rows = []
        seen_abbrs = set()

        for term_row in reducible_terms_df.to_dict("records"):
            term = str(term_row["term"]).strip()
            suggested_abbreviation = str(term_row["suggested_abbreviation"]).upper()

            # Если для этого сокращения есть канонический термин из глоссария,
            # используем его напрямую в merged.
            canonical_term = glossary_map.get(suggested_abbreviation)
            if canonical_term:
                term = canonical_term
                term_row["term"] = canonical_term
                term_row["normalized_term"] = self._normalize_term_for_compare(canonical_term)
                term_row["word_count"] = len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", canonical_term))

            matched_items = []
            for abbr_row in existing_records:
                found_abbreviation = str(abbr_row.get("abbreviation", ""))
                matched_term = str(abbr_row.get("matched_term", ""))

                if found_abbreviation.upper() == suggested_abbreviation or matched_term == term:
                    matched_items.append(abbr_row)

            if matched_items:
                best_match = max(
                    matched_items,
                    key=lambda x: (
                        self._source_priority(str(x.get("detection_type", ""))),
                        float(x.get("match_score", 0))
                    )
                )
                rows.append({
                    "term": term,
                    "normalized_term": term_row["normalized_term"],
                    "suggested_abbreviation": suggested_abbreviation,
                    "word_count": term_row["word_count"],
                    "frequency": term_row["frequency"],
                    "abbreviation_found_in_text": True,
                    "found_abbreviation": best_match.get("abbreviation", ""),
                    "match_score": best_match.get("match_score", 0),
                    "examples": term_row["examples"]
                })
            else:
                rows.append({
                    "term": term,
                    "normalized_term": term_row["normalized_term"],
                    "suggested_abbreviation": suggested_abbreviation,
                    "word_count": term_row["word_count"],
                    "frequency": term_row["frequency"],
                    "abbreviation_found_in_text": suggested_abbreviation in glossary_map,
                    "found_abbreviation": suggested_abbreviation if suggested_abbreviation in glossary_map else "",
                    "match_score": 100 if suggested_abbreviation in glossary_map else 0,
                    "examples": term_row["examples"]
                })

            seen_abbrs.add(suggested_abbreviation)

        # Гарантированно добавляем строки для всех глоссарных сокращений,
        # которых вдруг нет после этапа 1 и этапа 2.
        for abbr, canonical_term in glossary_map.items():
            if abbr in seen_abbrs:
                continue

            rows.append({
                "term": canonical_term,
                "normalized_term": self._normalize_term_for_compare(canonical_term),
                "suggested_abbreviation": abbr,
                "word_count": len(regex.findall(r"[A-Za-zА-Яа-яЁё-]+", canonical_term)),
                "frequency": 1,
                "abbreviation_found_in_text": True,
                "found_abbreviation": abbr,
                "match_score": 100,
                "examples": "Добавлено из раздела 'Обозначения и сокращения'"
            })

        df = pd.DataFrame(rows)
        if df.empty:
            return df

        # Схлопываем дубли по сокращению + канонической нормализованной форме.
        grouped = {}
        for row in df.to_dict("records"):
            key = (
                str(row.get("suggested_abbreviation", "")).upper(),
                str(row.get("normalized_term", "")).strip().lower()
            )
            if key not in grouped:
                grouped[key] = row
            else:
                grouped[key]["frequency"] = max(int(grouped[key].get("frequency", 0)), int(row.get("frequency", 0)))
                grouped[key]["abbreviation_found_in_text"] = bool(grouped[key].get("abbreviation_found_in_text", False)) or bool(row.get("abbreviation_found_in_text", False))
                grouped[key]["match_score"] = max(float(grouped[key].get("match_score", 0)), float(row.get("match_score", 0)))

        return pd.DataFrame(grouped.values()).sort_values(
            by=["abbreviation_found_in_text", "frequency", "word_count", "term"],
            ascending=[False, False, False, True]
        ).reset_index(drop=True)

    def _save_results(
        self,
        fragments_df: pd.DataFrame,
        reducible_terms_df: pd.DataFrame,
        existing_abbreviations_df: pd.DataFrame,
        merged_df: pd.DataFrame,
        output_dir: Path
    ) -> Dict[str, Path]:
        saved_files: Dict[str, Path] = {}

        fragments_csv = output_dir / "document_fragments.csv"
        fragments_xlsx = output_dir / "document_fragments.xlsx"
        reducible_csv = output_dir / "reducible_terms.csv"
        reducible_xlsx = output_dir / "reducible_terms.xlsx"
        existing_csv = output_dir / "existing_abbreviations.csv"
        existing_xlsx = output_dir / "existing_abbreviations.xlsx"
        merged_csv = output_dir / "merged_terms_and_abbreviations.csv"
        merged_xlsx = output_dir / "merged_terms_and_abbreviations.xlsx"

        fragments_df.to_csv(fragments_csv, index=False, encoding="utf-8-sig")
        reducible_terms_df.to_csv(reducible_csv, index=False, encoding="utf-8-sig")
        existing_abbreviations_df.to_csv(existing_csv, index=False, encoding="utf-8-sig")
        merged_df.to_csv(merged_csv, index=False, encoding="utf-8-sig")

        saved_files["document_fragments_csv"] = fragments_csv
        saved_files["reducible_terms_csv"] = reducible_csv
        saved_files["existing_abbreviations_csv"] = existing_csv
        saved_files["merged_csv"] = merged_csv

        try:
            fragments_df.to_excel(fragments_xlsx, index=False)
            reducible_terms_df.to_excel(reducible_xlsx, index=False)
            existing_abbreviations_df.to_excel(existing_xlsx, index=False)
            merged_df.to_excel(merged_xlsx, index=False)

            saved_files["document_fragments_xlsx"] = fragments_xlsx
            saved_files["reducible_terms_xlsx"] = reducible_xlsx
            saved_files["existing_abbreviations_xlsx"] = existing_xlsx
            saved_files["merged_xlsx"] = merged_xlsx
        except ModuleNotFoundError as exc:
            print("Внимание: не удалось сохранить XLSX-файлы.")
            print("Причина:", exc)
            print("CSV-файлы при этом успешно сохранены.")

        return saved_files


if __name__ == "__main__":
    input_docx = "input.docx"
    output_dir = "result_stage2"

    analyzer = Stage2ReductionAnalyzer()
    saved_files = analyzer.run(input_docx, output_dir)

    print("=" * 60)
    print("ЭТАП 2 ЗАВЕРШЁН: выделение словоформ и аббревиатур")
    print("=" * 60)
    print("Сформированы файлы:")
    for name, path in saved_files.items():
        print(f"{name}: {path}")
