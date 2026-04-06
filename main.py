from __future__ import annotations

# ------------------------------ #
# Импорт стандартных библиотек   #
# ------------------------------ #

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import List, Iterable, Tuple, Dict

# ------------------------------ #
# Импорт внешних библиотек       #
# ------------------------------ #

# python-docx:
# используется для чтения .docx-документа Microsoft Word
from docx import Document

# regex:
# более мощная альтернатива стандартному re;
# используем для извлечения слов из текста
import regex

# pandas:
# нужен для удобного представления и сохранения результатов анализа
import pandas as pd

# Natasha:
# в этой задаче используем её для сегментации текста на предложения
from natasha import Doc, Segmenter

# pymorphy2:
# морфологический анализатор русского языка;
# позволяет определить часть речи и нормальную форму слова
from pymorphy2 import MorphAnalyzer


# ============================================================
# 1. Служебные dataclass-структуры
# ============================================================

@dataclass
class TextFragment:
    """
    Фрагмент текста, извлечённый из документа.

    Поля:
    - source_type: откуда пришёл текст ('heading', 'paragraph', 'table_cell')
    - source_index: индекс фрагмента внутри своего типа
    - text: сам текст
    """
    source_type: str
    source_index: int
    text: str


@dataclass
class CandidateMention:
    """
    Конкретное вхождение словоформы / словосочетания,
    которое может быть сокращено.

    Это НЕ итоговая аббревиатура и НЕ решение о том,
    нужно ли её вводить. Это только кандидат,
    выделенный на этапе распознавания текста.

    Поля:
    - source_type: тип источника ('heading', 'paragraph', ...)
    - source_index: номер фрагмента
    - sentence: предложение, в котором найден кандидат
    - phrase: исходная словоформа / словосочетание
    - normalized_phrase: нормализованный вариант (леммы)
    - proposed_abbreviation: автоматически построенная аббревиатура
    - word_count: количество слов в кандидате
    """
    source_type: str
    source_index: int
    sentence: str
    phrase: str
    normalized_phrase: str
    proposed_abbreviation: str
    word_count: int


# ============================================================
# 2. Основной класс для распознавания текста
# ============================================================

class ReducibleWordformRecognizer:
    """
    Класс выполняет ПЕРВЫЙ этап проекта:
    распознавание текста и поиск словоформ / словосочетаний,
    которые потенциально можно сокращать.

    На текущем этапе класс решает именно эту подзадачу:
    1) читает документ .docx;
    2) извлекает текстовые фрагменты;
    3) делит текст на предложения;
    4) выполняет морфологический разбор слов;
    5) выделяет длинные терминоподобные конструкции,
       которые могут стать кандидатами на сокращение.

    Важно:
    - здесь мы ещё НЕ ищем уже существующие аббревиатуры;
    - здесь мы ещё НЕ определяем, надо ли вводить аббревиатуру;
    - здесь мы только готовим основу для следующих этапов.
    """

    def __init__(self) -> None:
        """
        Инициализация используемых NLP-инструментов.
        """
        # Segmenter из Natasha используется для разбиения текста на предложения
        self.segmenter = Segmenter()

        # MorphAnalyzer из pymorphy2 нужен для определения части речи
        # и получения нормальной формы слова
        self.morph = MorphAnalyzer()

        # Стоп-слова / слишком общие слова, которые сами по себе
        # не должны становиться основой для кандидата
        self.stop_lemmas = {
            "и", "или", "в", "во", "на", "по", "с", "со", "к", "ко",
            "у", "о", "об", "от", "до", "за", "из", "под", "над",
            "при", "для", "не", "это", "тот", "такой", "данный",
            "этот", "который", "как", "а", "но", "же"
        }

        # Разрешённые части речи для построения сокращаемых словосочетаний.
        # Здесь берём те POS, которые чаще всего формируют термины:
        # - NOUN  : существительное
        # - ADJF  : полное прилагательное
        # - ADJS  : краткое прилагательное
        # - PRTF  : полное причастие
        # - PRTS  : краткое причастие
        self.allowed_pos = {"NOUN", "ADJF", "ADJS", "PRTF", "PRTS"}

        # Регулярное выражение для выделения слов.
        # Берём русские и латинские буквы, допускаем дефис.
        self.word_pattern = regex.compile(r"[A-Za-zА-Яа-яЁё-]+")

    # ========================================================
    # 3. Работа с документом
    # ========================================================

    def load_docx_fragments(self, docx_path: str | Path) -> List[TextFragment]:
        """
        Читает .docx-файл и извлекает из него текстовые фрагменты.

        На текущем этапе извлекаются:
        - заголовки;
        - обычные абзацы;
        - текст из ячеек таблиц.

        Почему так:
        - для первого этапа нам важно начать с основного текстового потока;
        - поддержку сносок можно добавить отдельным шагом позже,
          если команда будет расширять модуль.

        Возвращает список TextFragment.
        """
        docx_path = Path(docx_path)
        doc = Document(docx_path)

        fragments: List[TextFragment] = []

        # --------- Извлечение абзацев и заголовков --------- #
        paragraph_index = 0
        heading_index = 0

        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text:
                continue

            # Пытаемся определить, является ли абзац заголовком
            style_name = ""
            if paragraph.style is not None and paragraph.style.name:
                style_name = paragraph.style.name.lower()

            if "heading" in style_name or "заголов" in style_name:
                fragments.append(
                    TextFragment(
                        source_type="heading",
                        source_index=heading_index,
                        text=text
                    )
                )
                heading_index += 1
            else:
                fragments.append(
                    TextFragment(
                        source_type="paragraph",
                        source_index=paragraph_index,
                        text=text
                    )
                )
                paragraph_index += 1

        # --------- Извлечение текста из таблиц --------- #
        table_cell_index = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._clean_text(cell.text)
                    if not text:
                        continue

                    fragments.append(
                        TextFragment(
                            source_type="table_cell",
                            source_index=table_cell_index,
                            text=text
                        )
                    )
                    table_cell_index += 1

        return fragments

    # ========================================================
    # 4. Предобработка текста
    # ========================================================

    def _clean_text(self, text: str) -> str:
        """
        Упрощённая очистка текста:
        - убираем лишние пробелы;
        - схлопываем повторяющиеся пробелы и переносы строк.

        Это нужно, чтобы одинаковые конструкции не отличались
        только количеством пробелов.
        """
        if not text:
            return ""

        text = text.replace("\n", " ")
        text = regex.sub(r"\s+", " ", text)
        return text.strip()

    def split_into_sentences(self, text: str) -> List[str]:
        """
        Делит текст на предложения.

        Используем Natasha, потому что она лучше подходит
        для русского текста, чем примитивное деление по точке.
        """
        if not text.strip():
            return []

        doc = Doc(text)
        doc.segment(self.segmenter)

        sentences = []
        for sent in doc.sents:
            sentence_text = sent.text.strip()
            if sentence_text:
                sentences.append(sentence_text)

        return sentences

    def extract_words(self, text: str) -> List[str]:
        """
        Извлекает слова из текста.

        Например:
        "Автоматизированная система мониторинга событий"
        ->
        ["Автоматизированная", "система", "мониторинга", "событий"]
        """
        return self.word_pattern.findall(text)

    # ========================================================
    # 5. Морфологический анализ
    # ========================================================

    def parse_word(self, word: str):
        """
        Выполняет морфологический разбор одного слова.

        Возвращает наиболее вероятный разбор из pymorphy2.
        """
        return self.morph.parse(word)[0]

    def get_normal_form(self, word: str) -> str:
        """
        Возвращает нормальную форму слова (лемму).

        Например:
        "событий" -> "событие"
        "автоматизированная" -> "автоматизированный"
        """
        return self.parse_word(word).normal_form

    def get_pos(self, word: str) -> str | None:
        """
        Возвращает часть речи слова.

        Например:
        - NOUN
        - ADJF
        - PRTF
        """
        return self.parse_word(word).tag.POS

    # ========================================================
    # 6. Правила отбора слов для потенциальных сокращений
    # ========================================================

    def is_content_word(self, word: str) -> bool:
        """
        Проверяет, можно ли считать слово содержательным
        для построения сокращаемой словоформы.

        Логика:
        - слово должно быть достаточно длинным;
        - его часть речи должна подходить;
        - лемма не должна быть стоп-словом.

        Это позволяет отсечь служебные слова:
        "и", "в", "на", "для" и т.п.
        """
        # Отсеиваем слишком короткие слова
        if len(word) < 3:
            return False

        parse = self.parse_word(word)
        lemma = parse.normal_form
        pos = parse.tag.POS

        if lemma in self.stop_lemmas:
            return False

        if pos not in self.allowed_pos:
            return False

        return True

    def is_term_like_chunk(self, chunk: List[Tuple[str, str, str]]) -> bool:
        """
        Проверяет, похож ли набор слов на термин / словосочетание,
        которое потенциально можно сокращать.

        Параметр chunk:
        список кортежей вида:
        (исходное_слово, лемма, часть_речи)

        Эвристика:
        - длина 2..6 слов;
        - последнее слово должно быть существительным;
        - в конструкции должно быть хотя бы одно существительное;
        - все слова должны быть содержательными.

        Почему именно так:
        - большинство сокращаемых терминов в технических текстах —
          это многословные именные группы:
          "автоматизированная система мониторинга событий",
          "центр обработки данных",
          "модуль анализа текста" и т.д.
        """
        if len(chunk) < 2 or len(chunk) > 6:
            return False

        pos_list = [pos for _, _, pos in chunk]

        # Последнее слово обычно является "головой" термина
        if pos_list[-1] != "NOUN":
            return False

        # В словосочетании должно быть хотя бы одно существительное
        if "NOUN" not in pos_list:
            return False

        return True

    def build_abbreviation(self, words: List[str]) -> str:
        """
        Строит предлагаемую аббревиатуру
        по первым буквам слов.

        Пример:
        "Автоматизированная система мониторинга событий"
        -> "АСМС"

        На первом этапе это вспомогательное поле.
        Оно нужно, чтобы потом на следующих этапах
        было проще анализировать перспективные кандидаты.
        """
        letters = []
        for word in words:
            clean_word = word.strip("-")
            if clean_word:
                letters.append(clean_word[0].upper())

        return "".join(letters)

    # ========================================================
    # 7. Выделение кандидатов из предложения
    # ========================================================

    def extract_candidates_from_sentence(
        self,
        sentence: str,
        source_type: str,
        source_index: int
    ) -> List[CandidateMention]:
        """
        Выделяет из одного предложения все терминоподобные конструкции,
        которые потенциально поддаются сокращению.

        Алгоритм:
        1. Извлекаем слова.
        2. Для каждого слова определяем:
           - нормальную форму,
           - часть речи.
        3. Оставляем только содержательные слова.
        4. Строим из них последовательности.
        5. Генерируем окна длиной 2..6 слов.
        6. Оставляем только те окна, которые похожи на термин.

        Возвращает список CandidateMention.
        """
        words = self.extract_words(sentence)

        # Если слов слишком мало, кандидатов быть не может
        if len(words) < 2:
            return []

        # Для каждого слова заранее рассчитываем лемму и POS
        analyzed_words: List[Tuple[str, str, str]] = []
        for word in words:
            lemma = self.get_normal_form(word)
            pos = self.get_pos(word)

            # Если часть речи не определилась, подставляем пустое значение
            if pos is None:
                pos = ""

            analyzed_words.append((word, lemma, pos))

        # Накапливаем "цепочки" подряд идущих содержательных слов.
        # Например, в предложении:
        # "В работе используется автоматизированная система мониторинга событий."
        # цепочка будет:
        # ["автоматизированная", "система", "мониторинга", "событий"]
        candidate_chains: List[List[Tuple[str, str, str]]] = []
        current_chain: List[Tuple[str, str, str]] = []

        for word, lemma, pos in analyzed_words:
            if self.is_content_word(word):
                current_chain.append((word, lemma, pos))
            else:
                if current_chain:
                    candidate_chains.append(current_chain)
                    current_chain = []

        # Не забываем последнюю цепочку, если предложение закончилось на ней
        if current_chain:
            candidate_chains.append(current_chain)

        results: List[CandidateMention] = []

        # Из каждой цепочки строим окна длиной от 2 до 6 слов
        for chain in candidate_chains:
            chain_candidates = self._extract_candidates_from_chain(
                chain=chain,
                sentence=sentence,
                source_type=source_type,
                source_index=source_index
            )
            results.extend(chain_candidates)

        return results

    def _extract_candidates_from_chain(
        self,
        chain: List[Tuple[str, str, str]],
        sentence: str,
        source_type: str,
        source_index: int
    ) -> List[CandidateMention]:
        """
        Вспомогательный метод:
        из одной цепочки содержательных слов генерирует
        все допустимые словосочетания-кандидаты.

        Пример цепочки:
        [
            ("Автоматизированная", "автоматизированный", "ADJF"),
            ("система", "система", "NOUN"),
            ("мониторинга", "мониторинг", "NOUN"),
            ("событий", "событие", "NOUN")
        ]

        Из неё могут быть получены окна:
        - "Автоматизированная система"
        - "система мониторинга"
        - "мониторинга событий"
        - "Автоматизированная система мониторинга"
        - "система мониторинга событий"
        - "Автоматизированная система мониторинга событий"
        """
        mentions: List[CandidateMention] = []

        # Ограничиваем максимальный размер окна,
        # чтобы не получать слишком длинные конструкции
        max_window = min(6, len(chain))

        for start in range(len(chain)):
            for window_size in range(2, max_window + 1):
                end = start + window_size
                if end > len(chain):
                    break

                chunk = chain[start:end]

                # Проверяем, похож ли этот кусок на термин
                if not self.is_term_like_chunk(chunk):
                    continue

                words = [word for word, _, _ in chunk]
                lemmas = [lemma for _, lemma, _ in chunk]

                phrase = " ".join(words)
                normalized_phrase = " ".join(lemmas)
                proposed_abbreviation = self.build_abbreviation(words)

                mentions.append(
                    CandidateMention(
                        source_type=source_type,
                        source_index=source_index,
                        sentence=sentence,
                        phrase=phrase,
                        normalized_phrase=normalized_phrase,
                        proposed_abbreviation=proposed_abbreviation,
                        word_count=len(words)
                    )
                )

        return mentions

    # ========================================================
    # 8. Обработка всего документа
    # ========================================================

    def analyze_document(self, docx_path: str | Path) -> List[CandidateMention]:
        """
        Основной метод первого этапа.

        Полный конвейер:
        1. читаем документ;
        2. разбиваем фрагменты на предложения;
        3. из каждого предложения извлекаем кандидатов;
        4. возвращаем общий список найденных кандидатов.

        Результат:
        список CandidateMention по всему документу.
        """
        fragments = self.load_docx_fragments(docx_path)

        all_mentions: List[CandidateMention] = []

        for fragment in fragments:
            sentences = self.split_into_sentences(fragment.text)

            for sentence in sentences:
                sentence_mentions = self.extract_candidates_from_sentence(
                    sentence=sentence,
                    source_type=fragment.source_type,
                    source_index=fragment.source_index
                )
                all_mentions.extend(sentence_mentions)

        return all_mentions

    # ========================================================
    # 9. Постобработка и агрегирование результатов
    # ========================================================

    def aggregate_mentions(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        """
        Преобразует список найденных вхождений в агрегированную таблицу.

        Зачем нужна агрегация:
        - одно и то же словосочетание может встретиться много раз;
        - для следующего этапа проекта удобнее видеть:
          что именно найдено, сколько раз встретилось,
          и какая аббревиатура предлагается.

        Возвращаем DataFrame со столбцами:
        - phrase_example
        - normalized_phrase
        - proposed_abbreviation
        - word_count
        - frequency
        - examples
        """
        if not mentions:
            return pd.DataFrame(columns=[
                "phrase_example",
                "normalized_phrase",
                "proposed_abbreviation",
                "word_count",
                "frequency",
                "examples"
            ])

        grouped: Dict[str, Dict] = {}

        for mention in mentions:
            key = mention.normalized_phrase

            if key not in grouped:
                grouped[key] = {
                    "phrase_example": mention.phrase,
                    "normalized_phrase": mention.normalized_phrase,
                    "proposed_abbreviation": mention.proposed_abbreviation,
                    "word_count": mention.word_count,
                    "frequency": 0,
                    "examples": []
                }

            grouped[key]["frequency"] += 1

            # Сохраняем не более 3 примеров предложений,
            # чтобы таблица не разрасталась слишком сильно
            if len(grouped[key]["examples"]) < 3:
                grouped[key]["examples"].append(mention.sentence)

        data = []
        for _, item in grouped.items():
            data.append({
                "phrase_example": item["phrase_example"],
                "normalized_phrase": item["normalized_phrase"],
                "proposed_abbreviation": item["proposed_abbreviation"],
                "word_count": item["word_count"],
                "frequency": item["frequency"],
                "examples": " || ".join(item["examples"])
            })

        df = pd.DataFrame(data)

        # Сортировка:
        # сначала по частоте (самые важные кандидаты наверх),
        # потом по длине словосочетания
        df = df.sort_values(
            by=["frequency", "word_count", "phrase_example"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

        return df

    def mentions_to_dataframe(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        """
        Преобразует "сырые" найденные вхождения в DataFrame.

        Это удобно, если нужно посмотреть каждое найденное место
        отдельно, а не только агрегированную статистику.
        """
        return pd.DataFrame([asdict(item) for item in mentions])

    # ========================================================
    # 10. Сохранение результатов
    # ========================================================

    def save_results(
        self,
        mentions: List[CandidateMention],
        output_dir: str | Path
    ) -> Dict[str, Path]:
        """
        Сохраняет результаты первого этапа в файлы.

        Создаёт:
        - raw_mentions.csv / xlsx  : все найденные вхождения;
        - aggregated_candidates.csv / xlsx : агрегированный список.

        Возвращает словарь с путями к сохранённым файлам.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        raw_df = self.mentions_to_dataframe(mentions)
        aggregated_df = self.aggregate_mentions(mentions)

        raw_csv = output_dir / "raw_mentions.csv"
        raw_xlsx = output_dir / "raw_mentions.xlsx"

        agg_csv = output_dir / "aggregated_candidates.csv"
        agg_xlsx = output_dir / "aggregated_candidates.xlsx"

        raw_df.to_csv(raw_csv, index=False, encoding="utf-8-sig")
        aggregated_df.to_csv(agg_csv, index=False, encoding="utf-8-sig")

        # Сохранение в Excel удобно для демонстрации куратору
        raw_df.to_excel(raw_xlsx, index=False)
        aggregated_df.to_excel(agg_xlsx, index=False)

        return {
            "raw_csv": raw_csv,
            "raw_xlsx": raw_xlsx,
            "aggregated_csv": agg_csv,
            "aggregated_xlsx": agg_xlsx
        }


# ============================================================
# 11. Точка входа для локального запуска
# ============================================================

if __name__ == "__main__":
    """
    Пример локального запуска.

    Что нужно сделать перед запуском:
    1. Положить .docx-файл рядом со скриптом
       или указать к нему полный путь.
    2. Установить зависимости:
       pip install python-docx regex pandas natasha pymorphy2 openpyxl

    Что получится на выходе:
    - папка result_stage1
    - таблицы с найденными словоформами / словосочетаниями,
      которые потенциально можно сокращать
    """

    input_docx = "test_reduction_input.docx"       # <-- замените на имя вашего файла
    output_dir = "result_stage1"

    recognizer = ReducibleWordformRecognizer()

    # Шаг 1. Анализируем документ и собираем все найденные кандидаты
    mentions = recognizer.analyze_document(input_docx)

    # Шаг 2. Сохраняем результат в файлы
    saved_files = recognizer.save_results(mentions, output_dir)

    # Шаг 3. Выводим краткую сводку в консоль
    print("=" * 60)
    print("ЭТАП 1 ЗАВЕРШЁН: распознавание текста и поиск кандидатов")
    print("=" * 60)
    print(f"Всего найдено вхождений-кандидатов: {len(mentions)}")
    print("Сформированы файлы:")

    for name, path in saved_files.items():
        print(f"{name}: {path}")