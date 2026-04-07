from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import List, Tuple, Dict

from docx import Document
import regex
import pandas as pd
from natasha import Doc, Segmenter
from pymorphy2 import MorphAnalyzer


@dataclass
class TextFragment:
    source_type: str
    source_index: int
    text: str


@dataclass
class CandidateMention:
    source_type: str
    source_index: int
    sentence: str
    phrase: str
    normalized_phrase: str
    proposed_abbreviation: str
    word_count: int


class ReducibleWordformRecognizer:
    """
    Этап 1 проекта:
    - чтение текста документа;
    - поиск терминоподобных словосочетаний;
    - выделение словоформ, которые потенциально поддаются сокращению.

    Версия v2 содержит более строгие фильтры качества:
    1) меньше перекрывающихся кандидатов;
    2) исключаются слишком общие словосочетания;
    3) в итог попадают только более "терминные" конструкции.
    """

    def __init__(self) -> None:
        self.segmenter = Segmenter()
        self.morph = MorphAnalyzer()
        self.word_pattern = regex.compile(r"[A-Za-zА-Яа-яЁё-]+")

        # Служебные/слишком общие леммы, которые часто создают шум.
        self.stop_lemmas = {
            "и", "или", "в", "во", "на", "по", "с", "со", "к", "ко",
            "у", "о", "об", "от", "до", "за", "из", "под", "над",
            "при", "для", "не", "это", "тот", "такой", "данный",
            "этот", "который", "как", "а", "но", "же", "либо"
        }

        # Очень общие термины-"мусор", которые часто встречаются в отчётах,
        # но редко имеют смысл как самостоятельные кандидаты на сокращение.
        self.generic_lemmas = {
            "задача", "результат", "пример", "случай", "этап", "часть",
            "область", "проект", "документ", "система", "средство",
            "анализ", "описание", "требование", "возможность", "критерий",
            "таблица", "рисунок", "данные", "подход", "метод", "работа",
            "процесс", "функция", "алгоритм", "решение", "способ",
            "порядок", "уровень", "качество", "контроль", "обработка"
        }

        self.allowed_pos = {"NOUN", "ADJF", "ADJS", "PRTF", "PRTS"}

    # -----------------------------
    # Чтение документа
    # -----------------------------

    def load_docx_fragments(self, docx_path: str | Path) -> List[TextFragment]:
        docx_path = Path(docx_path)
        doc = Document(docx_path)
        fragments: List[TextFragment] = []

        paragraph_index = 0
        heading_index = 0
        table_cell_index = 0

        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text:
                continue

            style_name = ""
            if paragraph.style is not None and paragraph.style.name:
                style_name = paragraph.style.name.lower()

            if "heading" in style_name or "заголов" in style_name:
                fragments.append(TextFragment("heading", heading_index, text))
                heading_index += 1
            else:
                fragments.append(TextFragment("paragraph", paragraph_index, text))
                paragraph_index += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._clean_text(cell.text)
                    if not text:
                        continue
                    fragments.append(TextFragment("table_cell", table_cell_index, text))
                    table_cell_index += 1

        return fragments

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\n", " ")
        text = regex.sub(r"\s+", " ", text)
        return text.strip()

    def split_into_sentences(self, text: str) -> List[str]:
        if not text.strip():
            return []
        doc = Doc(text)
        doc.segment(self.segmenter)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]

    def extract_words(self, text: str) -> List[str]:
        return self.word_pattern.findall(text)

    # -----------------------------
    # Морфология
    # -----------------------------

    def parse_word(self, word: str):
        return self.morph.parse(word)[0]

    def get_normal_form(self, word: str) -> str:
        return self.parse_word(word).normal_form

    def get_pos(self, word: str) -> str | None:
        return self.parse_word(word).tag.POS

    def is_content_word(self, word: str) -> bool:
        """
        Оставляем только содержательные слова, пригодные для терминов.
        """
        if len(word) < 3:
            return False

        # Полностью заглавные короткие слова обычно уже являются аббревиатурами,
        # а в этой задаче нам нужны именно ПОЛНЫЕ словоформы.
        if word.isupper() and len(word) <= 10:
            return False

        parse = self.parse_word(word)
        lemma = parse.normal_form
        pos = parse.tag.POS

        if lemma in self.stop_lemmas:
            return False
        if pos not in self.allowed_pos:
            return False
        return True

    # -----------------------------
    # Правила качества кандидатов
    # -----------------------------

    def is_term_like_chunk(self, chunk: List[Tuple[str, str, str]]) -> bool:
        """
        Более строгая проверка кандидата.

        Требования:
        - длина 2..5 слов;
        - последнее слово — существительное;
        - минимум 2 существительных ИЛИ одно существительное + один признак;
        - не допускаем слишком общие/мусорные конструкции.
        """
        if len(chunk) < 2 or len(chunk) > 5:
            return False

        words = [w for w, _, _ in chunk]
        lemmas = [l for _, l, _ in chunk]
        pos_list = [p for _, _, p in chunk]

        if pos_list[-1] != "NOUN":
            return False

        noun_count = sum(1 for p in pos_list if p == "NOUN")
        modifier_count = sum(1 for p in pos_list if p in {"ADJF", "ADJS", "PRTF", "PRTS"})

        # Конструкция должна быть достаточно "терминной".
        if not (noun_count >= 2 or (noun_count >= 1 and modifier_count >= 1)):
            return False

        # Отсекаем конструкции, состоящие только из слишком общих слов.
        meaningful_lemmas = [lemma for lemma in lemmas if lemma not in self.generic_lemmas]
        if len(meaningful_lemmas) < 2:
            return False

        # Если первый и последний элементы слишком общие, это обычно шум.
        if lemmas[0] in self.generic_lemmas and lemmas[-1] in self.generic_lemmas:
            return False

        # Минимальная суммарная "содержательность" по длине.
        if sum(len(w) for w in words) < 12:
            return False

        return True

    def build_abbreviation(self, words: List[str]) -> str:
        letters = []
        for word in words:
            clean_word = word.strip("-")
            if clean_word:
                letters.append(clean_word[0].upper())
        return "".join(letters)

    # -----------------------------
    # Выделение кандидатов
    # -----------------------------

    def extract_candidates_from_sentence(
        self,
        sentence: str,
        source_type: str,
        source_index: int,
    ) -> List[CandidateMention]:
        words = self.extract_words(sentence)
        if len(words) < 2:
            return []

        analyzed_words: List[Tuple[str, str, str]] = []
        for word in words:
            lemma = self.get_normal_form(word)
            pos = self.get_pos(word) or ""
            analyzed_words.append((word, lemma, pos))

        candidate_chains: List[List[Tuple[str, str, str]]] = []
        current_chain: List[Tuple[str, str, str]] = []

        for word, lemma, pos in analyzed_words:
            if self.is_content_word(word):
                current_chain.append((word, lemma, pos))
            else:
                if current_chain:
                    candidate_chains.append(current_chain)
                    current_chain = []

        if current_chain:
            candidate_chains.append(current_chain)

        results: List[CandidateMention] = []
        for chain in candidate_chains:
            results.extend(
                self._extract_best_candidates_from_chain(
                    chain=chain,
                    sentence=sentence,
                    source_type=source_type,
                    source_index=source_index,
                )
            )
        return results

    def _extract_best_candidates_from_chain(
        self,
        chain: List[Tuple[str, str, str]],
        sentence: str,
        source_type: str,
        source_index: int,
    ) -> List[CandidateMention]:
        """
        Главное отличие от первой версии:
        мы НЕ берём все возможные окна подряд.

        Сначала генерируем допустимые окна, затем удаляем подстроки,
        если внутри той же цепочки найден более длинный и более полный кандидат.
        Это сильно снижает шум вида:
        - "система мониторинга"
        - "мониторинга событий"
        - "автоматизированная система мониторинга"
        - ...
        и оставляет более полезные максимальные конструкции.
        """
        if len(chain) < 2:
            return []

        raw_windows = []
        max_window = min(5, len(chain))

        for start in range(len(chain)):
            for window_size in range(2, max_window + 1):
                end = start + window_size
                if end > len(chain):
                    break
                chunk = chain[start:end]
                if self.is_term_like_chunk(chunk):
                    raw_windows.append((start, end, chunk))

        if not raw_windows:
            return []

        # Оставляем только "максимальные" окна:
        # если текущее окно полностью содержится в более длинном, отбрасываем его.
        filtered_windows = []
        for i, (start_i, end_i, chunk_i) in enumerate(raw_windows):
            is_subwindow = False
            len_i = end_i - start_i

            for j, (start_j, end_j, chunk_j) in enumerate(raw_windows):
                if i == j:
                    continue

                len_j = end_j - start_j
                if start_j <= start_i and end_i <= end_j and len_j > len_i:
                    is_subwindow = True
                    break

            if not is_subwindow:
                filtered_windows.append((start_i, end_i, chunk_i))

        mentions: List[CandidateMention] = []
        for _, _, chunk in filtered_windows:
            words = [word for word, _, _ in chunk]
            lemmas = [lemma for _, lemma, _ in chunk]
            phrase = " ".join(words)
            normalized_phrase = " ".join(lemmas)
            abbreviation = self.build_abbreviation(words)

            mentions.append(
                CandidateMention(
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    phrase=phrase,
                    normalized_phrase=normalized_phrase,
                    proposed_abbreviation=abbreviation,
                    word_count=len(words),
                )
            )

        return mentions

    # -----------------------------
    # Анализ документа
    # -----------------------------

    def analyze_document(self, docx_path: str | Path) -> List[CandidateMention]:
        fragments = self.load_docx_fragments(docx_path)
        all_mentions: List[CandidateMention] = []

        for fragment in fragments:
            for sentence in self.split_into_sentences(fragment.text):
                all_mentions.extend(
                    self.extract_candidates_from_sentence(
                        sentence=sentence,
                        source_type=fragment.source_type,
                        source_index=fragment.source_index,
                    )
                )

        return all_mentions

    def mentions_to_dataframe(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        return pd.DataFrame([asdict(item) for item in mentions])

    def aggregate_mentions(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        if not mentions:
            return pd.DataFrame(columns=[
                "phrase_example",
                "normalized_phrase",
                "proposed_abbreviation",
                "word_count",
                "frequency",
                "examples",
            ])

        grouped: Dict[str, Dict] = {}
        for mention in mentions:
            key = mention.normalized_phrase
            if key not in grouped:
                grouped[key] = {
                    "phrase_example": mention.phrase,
                    "normalized_phrase": mention.normalized_phrase,
                    "proposed_abbreviation": mention.proposed_abbreviation,
                    "word_count": mention.word_count,
                    "frequency": 0,
                    "examples": [],
                }

            grouped[key]["frequency"] += 1
            if len(grouped[key]["examples"]) < 3:
                grouped[key]["examples"].append(mention.sentence)

        rows = []
        for item in grouped.values():
            rows.append({
                "phrase_example": item["phrase_example"],
                "normalized_phrase": item["normalized_phrase"],
                "proposed_abbreviation": item["proposed_abbreviation"],
                "word_count": item["word_count"],
                "frequency": item["frequency"],
                "examples": " || ".join(item["examples"]),
            })

        df = pd.DataFrame(rows)

        # Финальная очистка на уровне уже агрегированных кандидатов.
        # Логика:
        # - кандидаты из 2 слов оставляем только если они встречаются >= 2 раз;
        # - кандидаты из 3+ слов оставляем даже при 1 вхождении, так как они чаще терминные.
        df = df[
            ((df["word_count"] >= 3)) |
            ((df["word_count"] == 2) & (df["frequency"] >= 2))
        ].copy()

        df = df.sort_values(
            by=["frequency", "word_count", "phrase_example"],
            ascending=[False, False, True],
        ).reset_index(drop=True)

        return df

    def save_results(self, mentions: List[CandidateMention], output_dir: str | Path) -> Dict[str, Path]:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        raw_df = self.mentions_to_dataframe(mentions)
        aggregated_df = self.aggregate_mentions(mentions)

        raw_csv = output_dir / "raw_mentions.csv"
        raw_xlsx = output_dir / "raw_mentions.xlsx"
        agg_csv = output_dir / "aggregated_candidates.csv"
        agg_xlsx = output_dir / "aggregated_candidates.xlsx"

        raw_df.to_csv(raw_csv, index=False, encoding="utf-8-sig")
        aggregated_df.to_csv(agg_csv, index=False, encoding="utf-8-sig")
        raw_df.to_excel(raw_xlsx, index=False)
        aggregated_df.to_excel(agg_xlsx, index=False)

        return {
            "raw_csv": raw_csv,
            "raw_xlsx": raw_xlsx,
            "aggregated_csv": agg_csv,
            "aggregated_xlsx": agg_xlsx,
        }


if __name__ == "__main__":
    input_docx = "test_reduction_input.docx"
    output_dir = "result_stage1_v2"

    recognizer = ReducibleWordformRecognizer()
    mentions = recognizer.analyze_document(input_docx)
    saved_files = recognizer.save_results(mentions, output_dir)

    print("=" * 60)
    print("ЭТАП 1 (очищенная версия) ЗАВЕРШЁН")
    print("=" * 60)
    print(f"Всего найдено вхождений-кандидатов: {len(mentions)}")
    for name, path in saved_files.items():
        print(f"{name}: {path}")
