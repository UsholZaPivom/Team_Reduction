from __future__ import annotations

"""
text_recognition_candidates_v3.py

Этап 1 проекта:
распознавание текста документа и выделение словоформ / словосочетаний,
которые потенциально поддаются сокращению.

Дополнительные улучшения этой версии:
1. Убираются скобочные бытовые уточнения типа "(ноутбук)".
2. Жёстче фильтруются заголовочные конструкции:
   - "Технические решения ..."
   - "Решения по применению ..."
   - "Общие положения"
   - "Информационные активы"
   - "Описание объектов ..."
3. Не пропускаются фразы, начинающиеся с отглагольных форм:
   - применению ...
   - обеспечению ...
   - настройке ...
   - описанию ...
4. Длинные проектные названия режутся строже, чтобы не тащить
   "Реновация автоматизированной системы ..." как термин.
"""

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
import regex
import pandas as pd
from natasha import Doc, Segmenter
from pymorphy2 import MorphAnalyzer


@dataclass
class TextFragment:
    source_type: str
    source_index: int
    text: str


@dataclass
class CandidateMention:
    source_type: str
    source_index: int
    sentence: str
    phrase: str
    normalized_phrase: str
    proposed_abbreviation: str
    word_count: int


class ReducibleWordformRecognizerV3:
    def __init__(self) -> None:
        self.segmenter = Segmenter()
        self.morph = MorphAnalyzer()
        self.word_pattern = regex.compile(r"[A-Za-zА-Яа-яЁё-]+")

        self.allowed_pos = {"NOUN", "ADJF", "ADJS", "PRTF", "PRTS"}

        self.stop_lemmas = {
            "и", "или", "в", "во", "на", "по", "с", "со", "к", "ко",
            "у", "о", "об", "от", "до", "за", "из", "под", "над",
            "при", "для", "не", "это", "тот", "такой", "данный",
            "этот", "который", "как", "а", "но", "же"
        }

        self.generic_lemmas = {
            "задача", "результат", "пример", "этап", "область",
            "требование", "метод", "процесс", "алгоритм",
            "описание", "проект", "работа", "часть", "случай",
            "способ", "документ", "файл", "состав", "наименование",
            "решение", "объект", "уровень", "структура", "положение",
            "характеристика", "сведение", "параметр", "таблица",
            "рисунок", "приложение", "раздел", "подраздел",
            "функция", "мероприятие"
        }

        self.context_lemmas = {
            "отчёт", "документ", "тестирование", "проверка", "задача",
            "пример", "этап", "раздел", "предложение", "сочетание",
            "термин", "аббревиатура", "введение", "вывод", "анализ",
            "необходимость", "определение", "последующий", "существующий",
            "поздний", "короткий", "новый", "первый", "данный",
            "следующий", "описание", "проект", "работа", "содержание",
            "лист", "листов", "дата", "подпись", "рисунок", "таблица",
            "рамка", "текущий", "настоящий"
        }

        self.extra_noise_lemmas = {
            "рабочий", "группа", "ввод", "сокращение",
            "распознавание", "выделение", "основание", "лицензирование",
            "исполнитель", "заказчик", "площадка", "значение", "количество",
            "применение", "обеспечение", "описание", "настройка"
        }

        self.weak_head_lemmas = {
            "модуль", "блок", "часть", "уровень", "элемент", "подсистема",
            "состав", "объект", "средство", "система", "решение", "структура"
        }

        self.weak_left_noise_lemmas = {
            "состав", "наименование", "таблица", "рисунок", "приложение",
            "раздел", "подраздел", "параметр", "рамка", "проект",
            "технический", "решение", "общий", "настоящий",
            "применение", "обеспечение", "описание", "настройка",
            "функция", "мероприятие", "сведение", "характеристика"
        }

        self.weak_right_noise_lemmas = {
            "таблица", "рисунок", "раздел", "подраздел", "документ",
            "проект", "значение", "параметр"
        }

        self.service_line_markers = {
            "согласовано", "взам.", "инв.", "подп. и дата", "лист",
            "листов", "разраб.", "пров.", "н.контр.", "изм.", "кол.уч",
            "№ док", "подп.", "дата", "стадия"
        }

        self.heading_starts = (
            "общие положения",
            "описание объектов",
            "информационные активы",
            "технические решения",
            "решения по применению",
            "параметры настроек",
            "средство резервного копирования",
            "решения по применению встроенных",
            "решения по применению наложенных",
        )

        self.verbal_start_lemmas = {
            "применение", "обеспечение", "описание", "настройка",
            "использование", "реализация", "выполнение", "осуществление",
            "категорирование", "лицензирование"
        }

        self.project_noise_lemmas = {
            "реновация", "пояснительный", "записка", "организация", "ооо"
        }

    # ========================================================
    # Работа с документом
    # ========================================================

    def load_docx_fragments(self, docx_path: str | Path) -> List[TextFragment]:
        docx_path = Path(docx_path)
        doc = Document(docx_path)

        fragments: List[TextFragment] = []
        paragraph_index = 0
        heading_index = 0
        table_cell_index = 0

        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            if not text:
                continue

            style_name = ""
            if paragraph.style is not None and paragraph.style.name:
                style_name = paragraph.style.name.lower()

            if "heading" in style_name or "заголов" in style_name:
                fragments.append(TextFragment("heading", heading_index, text))
                heading_index += 1
            else:
                fragments.append(TextFragment("paragraph", paragraph_index, text))
                paragraph_index += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._clean_text(cell.text)
                    if not text:
                        continue
                    fragments.append(TextFragment("table_cell", table_cell_index, text))
                    table_cell_index += 1

        return fragments

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\n", " ")
        text = regex.sub(r"\s+", " ", text)
        return text.strip()

    def _remove_parenthetical_noise(self, text: str) -> str:
        """
        Убирает бытовые и уточняющие скобки, которые портят термин:
        инженерная станция (ноутбук) -> инженерная станция
        """
        if not text:
            return text

        def repl(match):
            inside = match.group(1).strip().lower()
            words = regex.findall(r"[A-Za-zА-Яа-яЁё-]+", inside)

            # Если внутри скобок короткое бытовое / уточняющее слово,
            # просто убираем скобки целиком.
            if len(words) <= 3:
                return " "

            return " " + match.group(0) + " "

        text = regex.sub(r"\(([^()]*)\)", repl, text)
        text = regex.sub(r"\s+", " ", text).strip()
        return text

    def split_into_sentences(self, text: str) -> List[str]:
        if not text.strip():
            return []
        doc = Doc(text)
        doc.segment(self.segmenter)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]

    def extract_words(self, text: str) -> List[str]:
        return self.word_pattern.findall(text)

    # ========================================================
    # Морфология
    # ========================================================

    def parse_word(self, word: str):
        return self.morph.parse(word)[0]

    def get_normal_form(self, word: str) -> str:
        return self.parse_word(word).normal_form

    def get_pos(self, word: str) -> str:
        pos = self.parse_word(word).tag.POS
        return pos if pos else ""

    # ========================================================
    # Фильтрация служебных фрагментов / предложений
    # ========================================================

    def is_existing_abbreviation_token(self, word: str) -> bool:
        clean_word = word.strip("-")
        if len(clean_word) < 2:
            return False
        if not regex.fullmatch(r"[A-Za-zА-Яа-яЁё-]+", clean_word):
            return False
        return clean_word.isupper() and len(clean_word) <= 10

    def should_skip_fragment(self, text: str) -> bool:
        t = self._clean_text(text)
        if not t:
            return True

        lowered = t.lower()

        if lowered in {"содержание", "обозначения и сокращения"}:
            return False

        if any(marker in lowered for marker in self.service_line_markers):
            return True

        if regex.fullmatch(r"(?:\d+(?:\.\d+)*)?\s*[А-ЯA-ZЁ].*?\d+\s*", t):
            if sum(ch.isdigit() for ch in t) >= 2 and len(t.split()) <= 8:
                return True

        if ("тел." in lowered or "факс" in lowered or "ул." in lowered or "договор от" in lowered):
            return True

        if regex.search(r"№\s*[\w/-]+", t) and len(t.split()) <= 12:
            return True

        if regex.search(r"\b\d{1,2}\.\d{1,2}\.\d{2,4}\b", t):
            return True

        digits = sum(ch.isdigit() for ch in t)
        letters = sum(ch.isalpha() for ch in t)
        if digits >= 6 and digits >= letters:
            return True

        if regex.search(r"\b(?:ssl|udp|порт|ip)\b", lowered) and digits >= 2:
            return True

        if regex.search(r"(?:ооо\s+«|инв\.|листов|стадия|подп\.)", lowered) and len(t.split()) <= 10:
            return True

        return False

    def is_service_sentence(self, sentence: str) -> bool:
        s = self._clean_text(sentence)
        if not s:
            return True

        lowered = s.lower()

        if self.should_skip_fragment(s):
            return True

        if lowered.startswith("рисунок") or lowered.startswith("таблица"):
            return True

        if ":" in s and sum(ch.isdigit() for ch in s) >= 2:
            return True

        if regex.fullmatch(r"(?:\d+(?:\.\d+)*)\s+.+", s) and len(s.split()) <= 6:
            return True

        if any(lowered.startswith(prefix) for prefix in self.heading_starts):
            return True

        return False

    # ========================================================
    # Фильтры слов
    # ========================================================

    def is_content_word(self, word: str) -> bool:
        if len(word) < 3:
            return False
        if self.is_existing_abbreviation_token(word):
            return False

        parse = self.parse_word(word)
        lemma = parse.normal_form
        pos = parse.tag.POS

        if lemma in self.stop_lemmas:
            return False
        if pos not in self.allowed_pos:
            return False

        return True

    # ========================================================
    # Фильтры словосочетаний
    # ========================================================

    def is_context_noise_phrase(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return True
        if lemmas[0] in self.context_lemmas:
            return True
        context_count = sum(1 for lemma in lemmas if lemma in self.context_lemmas)
        return context_count >= 2

    def is_generic_noise_phrase(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return True

        generic_count = sum(1 for lemma in lemmas if lemma in self.generic_lemmas)

        if len(lemmas) == 2 and generic_count >= 1:
            return True

        return generic_count >= max(2, len(lemmas) - 1)

    def is_process_noise_phrase(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return True
        noise_count = sum(1 for lemma in lemmas if lemma in self.extra_noise_lemmas)
        return noise_count >= 2

    def has_bad_boundaries(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return True
        if lemmas[0] in self.weak_left_noise_lemmas:
            return True
        if lemmas[-1] in self.weak_right_noise_lemmas:
            return True
        return False

    def is_incomplete_phrase(self, lemmas: List[str]) -> bool:
        if len(lemmas) != 2:
            return False
        if lemmas[0] in self.weak_head_lemmas:
            return True
        dependent_second_words = {
            "документ", "данные", "событие", "профиль", "защита",
            "информация", "оборудование", "управление", "решение"
        }
        return lemmas[1] in dependent_second_words

    def is_weak_document_phrase(self, lemmas: List[str]) -> bool:
        weak_pairs = {
            ("объект", "защита"),
            ("технический", "решение"),
            ("информационный", "актив"),
            ("средство", "защита"),
            ("состав", "уровень"),
            ("уровень", "управление"),
            ("механизм", "безопасность"),
            ("параметр", "установка"),
            ("параметр", "настройка"),
            ("общий", "положение"),
        }
        if len(lemmas) == 2 and tuple(lemmas) in weak_pairs:
            return True

        if len(lemmas) == 3 and lemmas[0] in {"состав", "уровень", "решение", "применение", "описание"}:
            return True

        return False

    def is_table_parameter_phrase(self, words: List[str], lemmas: List[str]) -> bool:
        raw = " ".join(words).lower()
        if regex.search(r"\b(?:ssl|udp|порт|сервер|ip|wifi|wi-fi)\b", raw):
            return True
        if any(lemma in {"значение", "параметр", "общий", "умолчание"} for lemma in lemmas):
            return True
        return False

    def starts_with_verbal_noise(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return True
        return lemmas[0] in self.verbal_start_lemmas

    def is_project_name_phrase(self, lemmas: List[str]) -> bool:
        if not lemmas:
            return False
        if lemmas[0] in self.project_noise_lemmas:
            return True
        if "реновация" in lemmas:
            return True
        if "ооо" in lemmas:
            return True
        return False

    def is_term_like_chunk(self, chunk: List[Tuple[str, str, str]]) -> bool:
        if len(chunk) < 2 or len(chunk) > 5:
            return False

        words = [word for word, _, _ in chunk]
        pos_list = [pos for _, _, pos in chunk]

        if pos_list[-1] != "NOUN":
            return False

        noun_count = sum(1 for pos in pos_list if pos == "NOUN")
        modifier_count = sum(1 for pos in pos_list if pos in {"ADJF", "ADJS", "PRTF", "PRTS"})

        if not (noun_count >= 2 or (noun_count >= 1 and modifier_count >= 1)):
            return False

        total_letters = sum(len(w) for w in words)
        if total_letters < 10:
            return False

        return True

    def build_abbreviation(self, words: List[str]) -> str:
        letters: List[str] = []
        for word in words:
            clean_word = word.strip("-")
            if clean_word:
                letters.append(clean_word[0].upper())
        return "".join(letters)

    # ========================================================
    # Выделение кандидатов
    # ========================================================

    def extract_candidates_from_sentence(self, sentence: str, source_type: str, source_index: int) -> List[CandidateMention]:
        sentence = self._remove_parenthetical_noise(sentence)

        if self.is_service_sentence(sentence):
            return []

        words = self.extract_words(sentence)
        if len(words) < 2:
            return []

        analyzed_words: List[Tuple[str, str, str]] = []
        for word in words:
            analyzed_words.append((word, self.get_normal_form(word), self.get_pos(word)))

        candidate_chains: List[List[Tuple[str, str, str]]] = []
        current_chain: List[Tuple[str, str, str]] = []

        for word, lemma, pos in analyzed_words:
            if self.is_content_word(word):
                current_chain.append((word, lemma, pos))
            else:
                if current_chain:
                    candidate_chains.append(current_chain)
                    current_chain = []

        if current_chain:
            candidate_chains.append(current_chain)

        all_mentions: List[CandidateMention] = []
        for chain in candidate_chains:
            all_mentions.extend(
                self._extract_candidates_from_chain(
                    chain=chain,
                    sentence=sentence,
                    source_type=source_type,
                    source_index=source_index
                )
            )

        return self.postfilter_sentence_mentions(all_mentions)

    def _extract_candidates_from_chain(
        self,
        chain: List[Tuple[str, str, str]],
        sentence: str,
        source_type: str,
        source_index: int
    ) -> List[CandidateMention]:
        candidates: List[Tuple[List[str], List[str], int, int]] = []
        max_window = min(5, len(chain))

        for start in range(len(chain)):
            for window_size in range(2, max_window + 1):
                end = start + window_size
                if end > len(chain):
                    break

                chunk = chain[start:end]
                if not self.is_term_like_chunk(chunk):
                    continue

                words = [word for word, _, _ in chunk]
                lemmas = [lemma for _, lemma, _ in chunk]

                if self.is_context_noise_phrase(lemmas):
                    continue
                if self.is_generic_noise_phrase(lemmas):
                    continue
                if self.is_process_noise_phrase(lemmas):
                    continue
                if self.is_incomplete_phrase(lemmas):
                    continue
                if self.has_bad_boundaries(lemmas):
                    continue
                if self.is_weak_document_phrase(lemmas):
                    continue
                if self.is_table_parameter_phrase(words, lemmas):
                    continue
                if self.starts_with_verbal_noise(lemmas):
                    continue
                if self.is_project_name_phrase(lemmas):
                    continue

                candidates.append((words, lemmas, start, end))

        if not candidates:
            return []

        maximal_candidates = self._keep_maximal_candidates(candidates)
        mentions: List[CandidateMention] = []

        for words, lemmas, _, _ in maximal_candidates:
            mentions.append(
                CandidateMention(
                    source_type=source_type,
                    source_index=source_index,
                    sentence=sentence,
                    phrase=" ".join(words),
                    normalized_phrase=" ".join(lemmas),
                    proposed_abbreviation=self.build_abbreviation(words),
                    word_count=len(words)
                )
            )

        return mentions

    def _keep_maximal_candidates(self, candidates: List[Tuple[List[str], List[str], int, int]]) -> List[Tuple[List[str], List[str], int, int]]:
        sorted_candidates = sorted(
            candidates,
            key=lambda item: (item[3] - item[2], len(" ".join(item[0]))),
            reverse=True
        )

        selected: List[Tuple[List[str], List[str], int, int]] = []
        for candidate in sorted_candidates:
            _, _, start, end = candidate
            is_nested = False
            for _, _, selected_start, selected_end in selected:
                if selected_start <= start and end <= selected_end:
                    is_nested = True
                    break
            if not is_nested:
                selected.append(candidate)

        return sorted(selected, key=lambda item: item[2])

    def postfilter_sentence_mentions(self, mentions: List[CandidateMention]) -> List[CandidateMention]:
        if not mentions:
            return mentions

        result: List[CandidateMention] = []
        for i, mention in enumerate(mentions):
            cur = mention.normalized_phrase.split()
            drop = False

            for j, other in enumerate(mentions):
                if i == j:
                    continue
                oth = other.normalized_phrase.split()
                if len(oth) == len(cur) + 1:
                    if oth[1:] == cur and oth[0] in self.weak_left_noise_lemmas:
                        drop = True
                        break
                    if oth[:-1] == cur and oth[-1] in self.weak_right_noise_lemmas:
                        drop = True
                        break

            if not drop:
                result.append(mention)

        return result

    # ========================================================
    # Обработка документа
    # ========================================================

    def analyze_document(self, docx_path: str | Path) -> List[CandidateMention]:
        fragments = self.load_docx_fragments(docx_path)
        all_mentions: List[CandidateMention] = []

        for fragment in fragments:
            cleaned_fragment_text = self._remove_parenthetical_noise(fragment.text)

            if self.should_skip_fragment(cleaned_fragment_text):
                continue

            for sentence in self.split_into_sentences(cleaned_fragment_text):
                all_mentions.extend(
                    self.extract_candidates_from_sentence(
                        sentence=sentence,
                        source_type=fragment.source_type,
                        source_index=fragment.source_index
                    )
                )

        return all_mentions

    # ========================================================
    # Агрегация и финальная очистка
    # ========================================================

    def mentions_to_dataframe(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        return pd.DataFrame([asdict(item) for item in mentions])

    def is_incomplete_two_word_phrase(self, lemmas: List[str], frequency: int = 1) -> bool:
        if len(lemmas) != 2:
            return False
        if lemmas[0] in self.weak_head_lemmas and frequency < 2:
            return True
        return False

    def postfilter_aggregated_candidates(self, df: pd.DataFrame) -> pd.DataFrame:
        if df.empty:
            return df

        rows = df.to_dict("records")
        filtered_rows = []
        rows_sorted = sorted(
            rows,
            key=lambda x: (x["word_count"], len(str(x["phrase_example"]))),
            reverse=True
        )

        kept_phrases: List[str] = []

        for row in rows_sorted:
            phrase = str(row["phrase_example"]).strip()
            normalized = str(row["normalized_phrase"]).strip().lower()
            lemmas = normalized.split()
            frequency = int(row.get("frequency", 1))
            lowered_phrase = phrase.lower()

            if self.has_bad_boundaries(lemmas):
                continue
            if self.is_incomplete_two_word_phrase(lemmas, frequency):
                continue
            if self.is_context_noise_phrase(lemmas):
                continue
            if self.is_generic_noise_phrase(lemmas):
                continue
            if self.is_process_noise_phrase(lemmas):
                continue
            if self.is_weak_document_phrase(lemmas):
                continue
            if self.starts_with_verbal_noise(lemmas):
                continue
            if self.is_project_name_phrase(lemmas):
                continue

            if regex.search(r"\b(?:таблица|рисунок|приложение|параметр|значение|сервер|порт)\b", lowered_phrase):
                continue

            if regex.search(r"\b(?:ооо|улица|российская|федерация|тел|факс)\b", lowered_phrase):
                continue

            if any(lowered_phrase.startswith(prefix) for prefix in self.heading_starts):
                continue

            is_subphrase = False
            for kept in kept_phrases:
                if lowered_phrase != kept and lowered_phrase in kept:
                    is_subphrase = True
                    break

            if is_subphrase:
                continue

            kept_phrases.append(lowered_phrase)
            filtered_rows.append(row)

        result = pd.DataFrame(filtered_rows)
        if result.empty:
            return result

        return result.sort_values(
            by=["frequency", "word_count", "phrase_example"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

    def aggregate_mentions(self, mentions: List[CandidateMention]) -> pd.DataFrame:
        if not mentions:
            return pd.DataFrame(columns=[
                "phrase_example",
                "normalized_phrase",
                "proposed_abbreviation",
                "word_count",
                "frequency",
                "examples"
            ])

        grouped: Dict[str, Dict] = {}
        for mention in mentions:
            key = mention.normalized_phrase
            if key not in grouped:
                grouped[key] = {
                    "phrase_example": mention.phrase,
                    "normalized_phrase": mention.normalized_phrase,
                    "proposed_abbreviation": mention.proposed_abbreviation,
                    "word_count": mention.word_count,
                    "frequency": 0,
                    "examples": []
                }

            grouped[key]["frequency"] += 1
            if len(grouped[key]["examples"]) < 3:
                grouped[key]["examples"].append(mention.sentence)

        rows = []
        for item in grouped.values():
            rows.append({
                "phrase_example": item["phrase_example"],
                "normalized_phrase": item["normalized_phrase"],
                "proposed_abbreviation": item["proposed_abbreviation"],
                "word_count": item["word_count"],
                "frequency": item["frequency"],
                "examples": " || ".join(item["examples"])
            })

        df = pd.DataFrame(rows)

        df = df[
            ((df["word_count"] == 2) & (df["frequency"] >= 2)) |
            ((df["word_count"] == 3) & (df["frequency"] >= 1)) |
            (df["word_count"] >= 4)
        ].copy()

        df = self.postfilter_aggregated_candidates(df)

        return df.sort_values(
            by=["frequency", "word_count", "phrase_example"],
            ascending=[False, False, True]
        ).reset_index(drop=True)

    # ========================================================
    # Сохранение
    # ========================================================

    def save_results(self, mentions: List[CandidateMention], output_dir: str | Path) -> Dict[str, Path]:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        raw_df = self.mentions_to_dataframe(mentions)
        agg_df = self.aggregate_mentions(mentions)

        raw_csv = output_dir / "raw_mentions.csv"
        agg_csv = output_dir / "aggregated_candidates.csv"
        raw_xlsx = output_dir / "raw_mentions.xlsx"
        agg_xlsx = output_dir / "aggregated_candidates.xlsx"

        raw_df.to_csv(raw_csv, index=False, encoding="utf-8-sig")
        agg_df.to_csv(agg_csv, index=False, encoding="utf-8-sig")

        saved_files: Dict[str, Path] = {
            "raw_csv": raw_csv,
            "aggregated_csv": agg_csv,
        }

        try:
            raw_df.to_excel(raw_xlsx, index=False)
            agg_df.to_excel(agg_xlsx, index=False)
            saved_files["raw_xlsx"] = raw_xlsx
            saved_files["aggregated_xlsx"] = agg_xlsx
        except ModuleNotFoundError as exc:
            print("Внимание: не удалось сохранить XLSX-файлы.")
            print("Причина:", exc)
            print("CSV-файлы при этом успешно сохранены.")

        return saved_files


if __name__ == "__main__":
    input_docx = "test_reduction_input.docx"
    output_dir = "result_stage1_v3"

    recognizer = ReducibleWordformRecognizerV3()
    mentions = recognizer.analyze_document(input_docx)
    saved_files = recognizer.save_results(mentions, output_dir)

    print("=" * 60)
    print("ЭТАП 1 V3 ЗАВЕРШЁН: распознавание текста и поиск кандидатов")
    print("=" * 60)
    print(f"Всего найдено вхождений-кандидатов: {len(mentions)}")
    print("Сформированы файлы:")
    for name, path in saved_files.items():
        print(f"{name}: {path}")
